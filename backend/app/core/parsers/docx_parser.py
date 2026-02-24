import asyncio
import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from app.core.parsers.base import BaseParser, ParsedContent

logger = logging.getLogger(__name__)

# Unsupported image formats (vector formats that XeLaTeX cannot handle directly)
_UNSUPPORTED_EXTENSIONS = {".wmf", ".emf"}


class DocxParser(BaseParser):
    async def parse(self, file_path: Path) -> ParsedContent:
        return await asyncio.to_thread(self._parse_sync, file_path)

    def _parse_sync(self, file_path: Path) -> ParsedContent:
        file_path = Path(file_path)
        doc = DocxDocument(str(file_path))

        full_text_parts: list[str] = []
        sections_list: list[dict] = []
        current_section: dict | None = None

        # --- Per-style formatting extraction ---
        style_formats: dict[str, dict] = {}

        # --- Image extraction state ---
        all_images: list[dict] = []
        image_counter = 0
        embed_to_filename: dict[str, str] = {}  # embed_id → filename for dedup

        all_paragraphs = doc.paragraphs
        para_idx = 0

        for child in doc.element.body:
            if child.tag == qn('w:p'):
                if para_idx < len(all_paragraphs):
                    para = all_paragraphs[para_idx]
                    para_idx += 1
                else:
                    continue
                text = para.text.strip()

                # Extract images from this paragraph
                para_images = self._extract_images_from_paragraph(
                    child, doc, image_counter, embed_to_filename,
                )
                image_counter += len([
                    img for img in para_images
                    if not img.get("skipped") and not img.get("dedup")
                ])

                # Skip only if both text and images are empty
                if not text and not para_images:
                    continue

                # Build paragraph output: text + image placeholders
                para_parts: list[str] = []
                if text:
                    para_parts.append(text)

                for img_info in para_images:
                    if img_info.get("skipped"):
                        para_parts.append(
                            f"% [IMAGE SKIPPED: unsupported format {img_info['ext']}]"
                        )
                    else:
                        width_str = f", width={img_info['width_cm']}cm" if img_info.get("width_cm") else ""
                        para_parts.append(
                            f"[IMAGE: {img_info['filename']}{width_str}]"
                        )
                        all_images.append(img_info)

                combined = "\n".join(para_parts)
                full_text_parts.append(combined)

                style_name = para.style.name if para.style else "Normal"

                if style_name not in style_formats and text:
                    style_formats[style_name] = self._extract_style_format(para)

                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    if current_section:
                        sections_list.append(current_section)
                    current_section = {"title": text, "content": ""}
                else:
                    if current_section is None:
                        current_section = {"title": "", "content": ""}
                    if current_section["content"]:
                        current_section["content"] += "\n"
                    current_section["content"] += combined

            elif child.tag == qn('w:tbl'):
                table_text = self._extract_table_as_text(child)
                if table_text:
                    full_text_parts.append(table_text)
                    if current_section is None:
                        current_section = {"title": "", "content": ""}
                    if current_section["content"]:
                        current_section["content"] += "\n"
                    current_section["content"] += table_text

        if current_section:
            sections_list.append(current_section)

        # --- Page layout per section ---
        page_layouts = []
        for i, section in enumerate(doc.sections):
            layout = {
                "section_index": i + 1,
                "page_width_cm": self._emu_to_cm(section.page_width),
                "page_height_cm": self._emu_to_cm(section.page_height),
                "left_margin_cm": self._emu_to_cm(section.left_margin),
                "right_margin_cm": self._emu_to_cm(section.right_margin),
                "top_margin_cm": self._emu_to_cm(section.top_margin),
                "bottom_margin_cm": self._emu_to_cm(section.bottom_margin),
                "header_distance_cm": self._emu_to_cm(section.header_distance),
                "footer_distance_cm": self._emu_to_cm(section.footer_distance),
            }
            # Header / footer text
            if section.header:
                h_text = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
                if h_text:
                    layout["header_text"] = h_text
            if section.footer:
                f_text = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
                if f_text:
                    layout["footer_text"] = f_text
            page_layouts.append(layout)

        # --- Tables ---
        tables_info = []
        for i, table in enumerate(doc.tables):
            t_info: dict = {
                "index": i + 1,
                "rows": len(table.rows),
                "cols": len(table.columns),
            }
            # First row as header sample
            if table.rows:
                t_info["header_cells"] = [cell.text.strip()[:40] for cell in table.rows[0].cells]
            # Check for merged cells
            merged = False
            for row in table.rows:
                seen = set()
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen:
                        merged = True
                        break
                    seen.add(cid)
                if merged:
                    break
            t_info["has_merged_cells"] = merged
            if len(table.rows) <= 30:
                content = []
                for row in table.rows:
                    row_cells = []
                    seen_tc = set()
                    for cell in row.cells:
                        cid = id(cell._tc)
                        if cid in seen_tc:
                            continue
                        seen_tc.add(cid)
                        row_cells.append(cell.text.strip()[:80])
                    content.append(row_cells)
                t_info["content"] = content
            tables_info.append(t_info)

        # --- Numbering / list formats ---
        numbering_info = self._extract_numbering(doc)

        # --- Document properties ---
        metadata: dict = {"filename": file_path.name}
        core = doc.core_properties
        if core.title:
            metadata["title"] = core.title
        if core.author:
            metadata["author"] = core.author

        # Build rich formatting metadata
        metadata["formatting"] = {
            "style_formats": style_formats,
            "page_layouts": page_layouts,
            "tables": tables_info,
            "numbering": numbering_info,
            "image_count": len(all_images),
        }

        return ParsedContent(
            text="\n\n".join(full_text_parts),
            metadata=metadata,
            sections=sections_list,
            images=all_images,
        )

    def _extract_images_from_paragraph(
        self,
        para_element,
        doc: DocxDocument,
        image_counter: int,
        embed_to_filename: dict[str, str],
    ) -> list[dict]:
        """Extract images from a paragraph's XML (w:drawing elements).

        Handles both wp:inline (inline images) and wp:anchor (floating images).
        Returns a list of image info dicts. Skipped formats get {"skipped": True}.
        """
        images: list[dict] = []
        # Namespace URIs
        ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
        ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        for drawing in para_element.iter(qn("w:drawing")):
            # Process both inline and anchor images
            for container_tag in (f"{{{ns_wp}}}inline", f"{{{ns_wp}}}anchor"):
                for container in drawing.iter(container_tag):
                    # Get dimensions from wp:extent
                    width_cm = None
                    height_cm = None
                    extent = container.find(f"{{{ns_wp}}}extent")
                    if extent is not None:
                        cx = extent.get("cx")
                        cy = extent.get("cy")
                        if cx:
                            width_cm = round(int(cx) / 360000, 1)
                        if cy:
                            height_cm = round(int(cy) / 360000, 1)

                    # Find the embed relationship ID (a:blip)
                    embed_id = None
                    for blip in container.iter(f"{{{ns_a}}}blip"):
                        embed_id = blip.get(f"{{{ns_r}}}embed")
                        if embed_id:
                            break

                    if not embed_id:
                        continue

                    # Check if this embed_id was already seen (dedup)
                    if embed_id in embed_to_filename:
                        existing_fn = embed_to_filename[embed_id]
                        # Find the existing image info to get its metadata
                        images.append({
                            "filename": existing_fn,
                            "width_cm": width_cm,
                            "height_cm": height_cm,
                            "dedup": True,
                        })
                        continue

                    # Get image data from relationship
                    try:
                        rel = doc.part.rels[embed_id]
                        target_part = rel.target_part
                        blob = target_part.blob
                        content_type = target_part.content_type
                    except (KeyError, AttributeError) as e:
                        logger.debug("Could not access image for embed_id=%s: %s", embed_id, e)
                        continue

                    ext = self._content_type_to_ext(content_type)

                    # Skip unsupported formats
                    if ext in _UNSUPPORTED_EXTENSIONS:
                        images.append({"skipped": True, "ext": ext})
                        continue

                    image_counter += 1
                    filename = f"figure_{image_counter:03d}{ext}"
                    embed_to_filename[embed_id] = filename

                    images.append({
                        "filename": filename,
                        "data": blob,
                        "content_type": content_type,
                        "width_cm": width_cm,
                        "height_cm": height_cm,
                    })

        return images

    @staticmethod
    def _content_type_to_ext(content_type: str) -> str:
        """Map MIME content type to file extension."""
        mapping = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
            "image/svg+xml": ".svg",
            "image/x-wmf": ".wmf",
            "image/x-emf": ".emf",
        }
        return mapping.get(content_type, ".png")

    @staticmethod
    def _extract_table_as_text(tbl_element) -> str:
        """Extract table content as formatted text from a w:tbl XML element."""
        rows_text = []
        for tr in tbl_element.findall(qn('w:tr')):
            cells = []
            for tc in tr.findall(qn('w:tc')):
                cell_texts = []
                for p in tc.findall(qn('w:p')):
                    t_parts = []
                    for t in p.iter(qn('w:t')):
                        if t.text:
                            t_parts.append(t.text)
                    cell_texts.append("".join(t_parts))
                cells.append(" ".join(cell_texts).strip()[:80])
            if any(cells):
                rows_text.append(" | ".join(cells))
        if not rows_text:
            return ""
        return "[表格]\n" + "\n".join(rows_text) + "\n[/表格]"

    def _extract_style_format(self, para) -> dict:
        """Extract detailed formatting for a paragraph's style + inline overrides."""
        info: dict = {}
        style = para.style

        # Walk the style inheritance chain to resolve font properties
        if style:
            self._resolve_style_fonts(style, info)

        # Run-level overrides (sample first non-empty run)
        for run in para.runs:
            if not run.text.strip():
                continue
            rf = run.font
            if rf.name and "font_ascii" not in info:
                info["font_ascii"] = rf.name
            ea = self._get_run_east_asian_font(run)
            if ea and "font_east_asian" not in info:
                info["font_east_asian"] = ea
            if rf.size and "font_size_pt" not in info:
                info["font_size_pt"] = rf.size.pt
            if rf.bold is not None and "bold" not in info:
                info["bold"] = rf.bold
            break

        # Paragraph formatting
        pf = para.paragraph_format
        if pf.alignment is not None:
            align_map = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY"}
            info["alignment"] = align_map.get(int(pf.alignment), str(pf.alignment))

        if pf.first_line_indent:
            info["first_line_indent_pt"] = round(pf.first_line_indent / 12700, 1)
        if pf.left_indent:
            info["left_indent_pt"] = round(pf.left_indent / 12700, 1)
        if pf.space_before:
            info["space_before_pt"] = round(pf.space_before / 12700, 1)
        if pf.space_after:
            info["space_after_pt"] = round(pf.space_after / 12700, 1)
        if pf.line_spacing:
            rule = int(pf.line_spacing_rule) if pf.line_spacing_rule is not None else None
            # AT_LEAST=3, EXACTLY=4: value is EMU (Length), convert to pt
            if rule in (3, 4):
                info["line_spacing_pt"] = round(pf.line_spacing / 12700, 1)
                info["line_spacing_rule"] = "EXACTLY" if rule == 4 else "AT_LEAST"
            elif rule in (0, 1, 2):
                # SINGLE=0, ONE_POINT_FIVE=1, DOUBLE=2: predefined proportions
                prop_map = {0: 1.0, 1: 1.5, 2: 2.0}
                info["line_spacing"] = prop_map[rule]
            else:
                # MULTIPLE=5 or unknown: float proportion
                info["line_spacing"] = float(pf.line_spacing)

        info["sample_text"] = para.text.strip()[:60]
        return info

    def _resolve_style_fonts(self, style, info: dict, depth: int = 0) -> None:
        """Walk the style inheritance chain to resolve font/bold/italic/size."""
        if depth > 10:
            return

        if style.font:
            sf = style.font
            if sf.name and "font_ascii" not in info:
                info["font_ascii"] = sf.name
            ea_font = self._get_east_asian_font(style)
            if ea_font and "font_east_asian" not in info:
                info["font_east_asian"] = ea_font
            if sf.size and "font_size_pt" not in info:
                info["font_size_pt"] = sf.size.pt
            if sf.bold is not None and "bold" not in info:
                info["bold"] = sf.bold
            if sf.italic is not None and "italic" not in info:
                info["italic"] = sf.italic

        # If still missing properties, check base (parent) style
        missing = {"font_ascii", "font_east_asian", "font_size_pt", "bold"} - set(info.keys())
        if missing and style.base_style:
            self._resolve_style_fonts(style.base_style, info, depth + 1)

    def _extract_numbering(self, doc) -> list[dict]:
        """Extract heading numbering formats from document XML."""
        results = []
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return results

        # Find heading-related numbering definitions
        numbering_xml = numbering_part._element
        for abstract_num in numbering_xml.findall(qn("w:abstractNum")):
            for lvl in abstract_num.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl"), "")
                num_fmt_elem = lvl.find(qn("w:numFmt"))
                lvl_text_elem = lvl.find(qn("w:lvlText"))
                if num_fmt_elem is not None and lvl_text_elem is not None:
                    fmt = num_fmt_elem.get(qn("w:val"), "")
                    text = lvl_text_elem.get(qn("w:val"), "")
                    if fmt == "decimal" and "%" in text:
                        results.append({
                            "level": int(ilvl),
                            "format": fmt,
                            "pattern": text,
                        })
        # Deduplicate by level, keep first
        seen = set()
        unique = []
        for item in results:
            if item["level"] not in seen:
                seen.add(item["level"])
                unique.append(item)
        return sorted(unique, key=lambda x: x["level"])[:6]

    @staticmethod
    def _get_east_asian_font(style) -> str | None:
        """Get East Asian font name from style XML."""
        try:
            rpr = style.element.find(qn("w:rPr"))
            if rpr is not None:
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    return rfonts.get(qn("w:eastAsia"))
        except Exception:
            pass
        return None

    @staticmethod
    def _get_run_east_asian_font(run) -> str | None:
        """Get East Asian font name from run XML."""
        try:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    return rfonts.get(qn("w:eastAsia"))
        except Exception:
            pass
        return None

    @staticmethod
    def _emu_to_cm(emu_value) -> float | None:
        """Convert EMU (English Metric Units) to centimeters."""
        if emu_value is None:
            return None
        return round(emu_value / 360000, 2)
