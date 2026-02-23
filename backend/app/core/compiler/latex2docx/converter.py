"""Main LaTeX-to-DOCX converter.

Consumes a token stream and builds a python-docx Document, dispatching
to environment and command handlers via a registry.
"""

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Emu, RGBColor

from app.core.compiler.word_preprocessor import WordExportMetadata
from app.core.compiler.word_postprocessor import (
    _fix_styles as fix_styles,
    _fix_page_layout as fix_page_layout,
    _add_page_numbers as add_page_numbers,
    _enable_update_fields as enable_update_fields,
    _parse_length as parse_length,
    _make_page_break as make_page_break,
    _make_section_break as make_section_break,
    _make_toc_field_paragraph as make_toc_field_paragraph,
)
from .profile import DocxProfile
from .tex_auxfiles import TexStructure
from .tokenizer import Token, TokenType, tokenize
from .text_utils import normalize_latex_text as _normalize_latex_text, SYMBOL_MAP as _SYMBOL_MAP

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Format state for nested inline formatting
# ---------------------------------------------------------------------------

@dataclass
class FormatState:
    """Tracks current inline formatting."""
    bold: bool = False
    italic: bool = False
    underline: bool = False
    font_name: str | None = None  # CJK font override
    font_size: int | None = None  # in half-points
    superscript: bool = False
    subscript: bool = False
    color: str | None = None      # hex RGB e.g. "FF0000"
    url: str | None = None        # for hyperlinks

    def copy(self) -> "FormatState":
        return FormatState(
            bold=self.bold, italic=self.italic, underline=self.underline,
            font_name=self.font_name, font_size=self.font_size,
            superscript=self.superscript, subscript=self.subscript,
            color=self.color, url=self.url,
        )

    def merge(self, **kwargs) -> "FormatState":
        new = self.copy()
        for k, v in kwargs.items():
            setattr(new, k, v)
        return new


# ---------------------------------------------------------------------------
# Section counters for heading numbering
# ---------------------------------------------------------------------------

@dataclass
class SectionCounters:
    """Maintain chapter/section/subsection numbering."""
    chapter: int = 0
    section: int = 0
    subsection: int = 0
    subsubsection: int = 0
    profile: DocxProfile | None = None

    @property
    def _unnumbered(self) -> set[str]:
        if self.profile:
            return set(self.profile.numbering.unnumbered_headings)
        return {"摘要", "abstract", "Abstract", "ABSTRACT", "致谢",
                "参考文献", "附录", "目录", "目  录", "References"}

    def increment(self, level: int) -> str | None:
        """Increment counter for *level* and return the formatted number string.

        level 1 = chapter, 2 = section, 3 = subsection, 4 = subsubsection
        Returns None if the heading should not be numbered.
        """
        if level == 1:
            self.chapter += 1
            self.section = 0
            self.subsection = 0
            self.subsubsection = 0
            return None  # chapter numbering handled specially
        elif level == 2:
            self.section += 1
            self.subsection = 0
            self.subsubsection = 0
            return f"{self.chapter}.{self.section}"
        elif level == 3:
            self.subsection += 1
            self.subsubsection = 0
            return f"{self.chapter}.{self.section}.{self.subsection}"
        elif level == 4:
            self.subsubsection += 1
            return f"{self.chapter}.{self.section}.{self.subsection}.{self.subsubsection}"
        return None

    def format_chapter(self, title: str) -> str:
        if title in self._unnumbered:
            return title
        if self.profile:
            return self.profile.format_chapter(self.chapter, title)
        return f"第 {self.chapter} 章  {title}"


# ---------------------------------------------------------------------------
# Main converter class
# ---------------------------------------------------------------------------

class LatexToDocxConverter:
    """Converts LaTeX token stream into a python-docx Document."""

    def __init__(
        self,
        metadata: WordExportMetadata | None = None,
        template_id: str = "",
        image_base_dir: Path | None = None,
        profile: DocxProfile | None = None,
        doc_class_type: str = "report",
        tex_structure: TexStructure | None = None,
    ):
        self.metadata = metadata or WordExportMetadata()
        self.template_id = template_id
        self.image_base_dir = image_base_dir or Path(".")
        self.profile = profile or DocxProfile()
        self.doc_class_type = self.profile.doc_class_type or doc_class_type
        self.tex_structure = tex_structure

        self.doc = Document()
        self.tokens: list[Token] = []
        self.pos = 0

        self.counters = SectionCounters(profile=self.profile)
        self.format_stack: list[FormatState] = [FormatState()]

        # Current paragraph being built (None means we need a new one)
        self._current_para = None
        self._para_alignment = None

        # Track whether we are in the document body
        self._in_document = False
        # Environment stack
        self._env_stack: list[str] = []

        # Caption/label state for floats
        self._pending_caption: str | None = None
        self._pending_label: str | None = None

        # Figure/table/equation counters
        self._figure_count = 0
        self._table_count = 0
        self._equation_count = 0
        self._footnote_count = 0
        self._footnotes: list[tuple[int, str]] = []
        # Flag: just exited a float → next paragraph gets extra space_before
        self._after_float = False

        # List nesting depth
        self._list_depth = 0
        self._list_type_stack: list[str] = []  # "itemize" or "enumerate"
        self._list_counters: list[int] = []     # for enumerate

        # Commands that take no arguments and should be silently skipped
        self._SKIP_COMMANDS = {
            "centering", "noindent", "raggedright", "raggedleft",
            "normalfont", "selectfont", "bfseries", "itshape",
            "upshape", "mdseries", "rmfamily", "sffamily", "ttfamily",
            "normalsize", "small", "footnotesize", "scriptsize", "tiny",
            "large", "Large", "LARGE", "huge", "Huge",
            "frontmatter", "mainmatter", "backmatter",
            "onehalfspacing", "doublespacing", "singlespacing",
            "protect", "relax", "indent", "par",
            "bigskip", "medskip", "smallskip",
            "vfill", "hfill", "dotfill", "hrulefill",
            "strut", "null", "phantom",
            "maketitle", "MAKETITLE", "makedeclaration",
            "appendix",
        }

        # Commands that take one {} arg but should be skipped entirely
        self._SKIP_WITH_ARG = {
            "pagestyle", "thispagestyle", "pagenumbering",
            "setlength", "addtolength", "setcounter", "addtocounter",
            "linespread", "bibliographystyle",
            "hypersetup", "captionsetup",
            "renewcommand", "newcommand", "providecommand",
            "DeclareCaptionFont",
            "newgeometry", "restoregeometry",
            "fancyhead", "fancyfoot", "fancyhf",
            "titleformat", "titlespacing", "titlecontents",
            "intobmk",
        }

    # ── Token access helpers ──────────────────────────────────────────

    def _peek(self, offset: int = 0) -> Token:
        idx = self.pos + offset
        if idx < len(self.tokens):
            return self.tokens[idx]
        return Token(type=TokenType.EOF, value="", pos=-1)

    def _advance(self) -> Token:
        tok = self._peek()
        self.pos += 1
        return tok

    def _skip_whitespace(self):
        while self._peek().type == TokenType.WHITESPACE:
            self.pos += 1

    def _expect(self, tok_type: TokenType) -> Token:
        tok = self._advance()
        if tok.type != tok_type:
            logger.debug("Expected %s, got %s at pos %d", tok_type, tok.type, tok.pos)
        return tok

    def _read_brace_group(self) -> str:
        """Read a {}-delimited group and return the text content."""
        self._skip_whitespace()
        if self._peek().type != TokenType.BRACE_OPEN:
            return ""
        self._advance()  # consume {
        depth = 1
        parts = []
        while depth > 0 and self._peek().type != TokenType.EOF:
            tok = self._advance()
            if tok.type == TokenType.BRACE_OPEN:
                depth += 1
                parts.append("{")
            elif tok.type == TokenType.BRACE_CLOSE:
                depth -= 1
                if depth > 0:
                    parts.append("}")
            elif tok.type == TokenType.COMMAND:
                parts.append(f"\\{tok.extra.get('name', '')}")
            else:
                parts.append(tok.value)
        return "".join(parts)

    def _read_brace_group_tokens(self) -> list[Token]:
        """Read a {}-delimited group and return the tokens inside."""
        self._skip_whitespace()
        if self._peek().type != TokenType.BRACE_OPEN:
            return []
        self._advance()  # consume {
        depth = 1
        result = []
        while depth > 0 and self._peek().type != TokenType.EOF:
            tok = self._advance()
            if tok.type == TokenType.BRACE_OPEN:
                depth += 1
                result.append(tok)
            elif tok.type == TokenType.BRACE_CLOSE:
                depth -= 1
                if depth > 0:
                    result.append(tok)
            else:
                result.append(tok)
        return result

    def _read_optional_arg(self) -> str | None:
        """Read an optional [...] argument, or return None."""
        self._skip_whitespace()
        if self._peek().type != TokenType.BRACKET_OPEN:
            return None
        self._advance()  # consume [
        parts = []
        depth = 1
        while depth > 0 and self._peek().type != TokenType.EOF:
            tok = self._advance()
            if tok.type == TokenType.BRACKET_OPEN:
                depth += 1
                parts.append("[")
            elif tok.type == TokenType.BRACKET_CLOSE:
                depth -= 1
                if depth > 0:
                    parts.append("]")
            else:
                parts.append(tok.value)
        return "".join(parts)

    # ── Paragraph management ─────────────────────────────────────────

    def _ensure_paragraph(self):
        """Return the current paragraph, creating one if needed."""
        if self._current_para is None:
            self._current_para = self.doc.add_paragraph()
            if self._para_alignment is not None:
                self._current_para.alignment = self._para_alignment
            # Apply list indent if we're in a list (includes hanging indent)
            if self._list_depth > 0:
                self._apply_list_indent(self._current_para)
            elif self._para_alignment is not None:
                # Remove first-line indent for centered/aligned paragraphs
                self._current_para.paragraph_format.first_line_indent = Pt(0)
            else:
                # Normal body paragraph — restore first-line indent from profile
                indent_pt = self.profile.styles.normal.first_line_indent_pt
                if indent_pt:
                    self._current_para.paragraph_format.first_line_indent = Pt(indent_pt)
            # Extra space after floats (like LaTeX's \textfloatsep ≈ 12pt)
            if self._after_float:
                self._current_para.paragraph_format.space_before = Pt(12)
                self._after_float = False
        return self._current_para

    def _finish_paragraph(self):
        """End the current paragraph so the next text starts a new one."""
        self._current_para = None
        self._para_alignment = None

    def _add_run(self, text: str, fmt: FormatState | None = None):
        """Add a text run to the current paragraph with formatting."""
        if not text:
            return
        text = _normalize_latex_text(text)
        para = self._ensure_paragraph()
        run = para.add_run(text)

        fmt = fmt or self.format_stack[-1]

        if fmt.bold:
            run.bold = True
        if fmt.italic:
            run.italic = True
        if fmt.underline:
            run.underline = True
        if fmt.superscript:
            run.font.superscript = True
        if fmt.subscript:
            run.font.subscript = True

        # Font
        run.font.name = self.profile.fonts.body_latin
        if fmt.font_name:
            run.font.name = fmt.font_name
            # Also set East Asian font
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), fmt.font_name)

        if fmt.font_size:
            run.font.size = Pt(fmt.font_size / 2)

        if fmt.color:
            run.font.color.rgb = RGBColor.from_string(fmt.color)

    def _add_hyperlink(self, url: str, display_text: str):
        """Add a hyperlink to the current paragraph."""
        para = self._ensure_paragraph()
        # Create hyperlink element
        part = para.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        run.append(rPr)

        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = display_text
        run.append(t)

        hyperlink.append(run)
        para._element.append(hyperlink)

    # ── List management ──────────────────────────────────────────────

    def _apply_list_indent(self, para):
        """Apply indentation for list items based on nesting depth.

        Uses a hanging-indent layout so the bullet/number hangs to the
        left of the wrapped text, matching LaTeX's itemize/enumerate.
        """
        # LaTeX itemize: wrapped text aligns with body text (left margin),
        # bullet is indented on the first line via positive first_line_indent.
        pf = para.paragraph_format
        pf.left_indent = Pt(0)
        pf.first_line_indent = Cm(0.55 * self._list_depth)

    # ── Heading ──────────────────────────────────────────────────────

    def _add_heading(self, level: int, starred: bool = False):
        """Add a heading. Reads the title from the next brace group."""
        title_tokens = self._read_brace_group_tokens()
        title = self._tokens_to_text(title_tokens)

        # Determine if numbered
        if starred or title in self.counters._unnumbered:
            display_title = title
        else:
            # Map Word heading level → TeX sectioning name for .aux lookup
            if self.doc_class_type in ("report", "book"):
                _level_names = {1: "chapter", 2: "section", 3: "subsection", 4: "subsubsection"}
            else:
                _level_names = {1: "section", 2: "subsection", 3: "subsubsection"}

            # Try TeX .aux data first
            tex_entry = None
            if self.tex_structure:
                level_name = _level_names.get(level)
                if level_name:
                    tex_entry = self.tex_structure.find_heading(title, level_name)

            if tex_entry:
                display_title = tex_entry.full_title
                # Still increment counters for fallback consistency
                self.counters.increment(level)
            elif level == 1 and self.doc_class_type in ("report", "book"):
                self.counters.increment(1)
                display_title = self.counters.format_chapter(title)
            elif level >= 2:
                num = self.counters.increment(level)
                if num:
                    display_title = self.profile.format_section(
                        level, title,
                        chapter=self.counters.chapter,
                        section=self.counters.section,
                        subsection=self.counters.subsection,
                        subsubsection=self.counters.subsubsection,
                    )
                else:
                    display_title = title
            else:
                # article-type: section is level 1
                num = self.counters.increment(level)
                if num:
                    display_title = self.profile.format_section(
                        level, title,
                        chapter=self.counters.chapter,
                        section=self.counters.section,
                        subsection=self.counters.subsection,
                        subsubsection=self.counters.subsubsection,
                    )
                else:
                    display_title = title

        # Reset per-chapter counters (equation, figure, table for per-chapter numbering)
        if level == 1 and not starred:
            self._equation_count = 0

        exclude_from_toc = starred or title in self.counters._unnumbered

        self._finish_paragraph()

        # Numbered body headings use built-in Heading styles so Word can
        # recognize "标题 1/2/3" in the style gallery and navigation.
        # Front-matter/starred headings remain custom to avoid TOC pollution.
        if exclude_from_toc:
            self._add_heading_no_toc(display_title, level)
        else:
            self._add_heading_builtin(display_title, level)

        self._finish_paragraph()

    def _add_heading_builtin(self, title: str, level: int):
        """Add a built-in Heading paragraph (Heading 1..4)."""
        style_name = f"Heading {min(max(level, 1), 4)}"
        try:
            para = self.doc.add_paragraph(style=style_name)
        except KeyError:
            para = self.doc.add_paragraph(style="Heading 1")

        # Defensive cleanup: remove direct numPr to avoid phantom bullets.
        pPr = para._element.get_or_add_pPr()
        for numPr in pPr.findall(qn("w:numPr")):
            pPr.remove(numPr)

        run = para.add_run(title)
        hs = self.profile.get_heading_style(level)
        if hs:
            run.font.size = Pt(hs.font_size_pt)
            run.bold = hs.bold
        else:
            run.font.size = Pt(15)
            run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        if self.profile.fonts.heading_east_asian:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), self.profile.fonts.heading_east_asian)

        if level == 1:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return para

    def _ensure_latex_heading_styles(self):
        """Create custom ``LaTeX Heading N`` styles (1-4) if they don't exist.

        Custom styles avoid phantom dots/bullets that some Word versions
        add to built-in ``Heading N`` styles via Normal.dotm numbering.
        """
        if getattr(self, "_heading_styles_created", False):
            return
        self._heading_styles_created = True

        for level in range(1, 5):
            style_name = f"LaTeX Heading {level}"
            try:
                self.doc.styles[style_name]
                continue  # already exists
            except KeyError:
                pass
            style = self.doc.styles.add_style(style_name, 1)  # 1 = paragraph
            style.base_style = self.doc.styles["Normal"]
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.name = self.profile.fonts.heading_latin
            if self.profile.fonts.heading_east_asian:
                _rPr = style.element.get_or_add_rPr()
                _rFonts = _rPr.find(qn("w:rFonts"))
                if _rFonts is None:
                    _rFonts = OxmlElement("w:rFonts")
                    _rPr.insert(0, _rFonts)
                _rFonts.set(qn("w:eastAsia"), self.profile.fonts.heading_east_asian)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.first_line_indent = Pt(0)
            style.paragraph_format.keep_with_next = True

            hs = self.profile.get_heading_style(level)
            if hs:
                style.font.size = Pt(hs.font_size_pt)

    def _add_heading_no_toc(self, title: str, level: int):
        """Add a heading that visually matches Heading N but is excluded from TOC.

        Uses a custom ``LaTeX Heading N`` style that Word's Normal.dotm
        won't auto-number, and that ``TOC \\o`` won't pick up (no outlineLvl).
        """
        self._ensure_latex_heading_styles()

        style_name = f"LaTeX Heading {min(level, 4)}"
        para = self.doc.add_paragraph(style=style_name)

        run = para.add_run(title)
        # Override font size per level (style has default)
        hs = self.profile.get_heading_style(level)
        if hs:
            run.font.size = Pt(hs.font_size_pt)
            run.bold = hs.bold
        else:
            run.font.size = Pt(15)
            run.bold = True
        # Explicit black — don't rely on style inheritance
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Center chapter-level headings (level 1) to match LaTeX PDF
        if level == 1:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return para

    def _tokens_to_text(self, tokens: list[Token]) -> str:
        """Convert tokens to readable text with basic ref/cite resolution."""
        from .math_handler import _latex_math_to_text

        def _skip_ws(idx: int) -> int:
            while idx < len(tokens) and tokens[idx].type == TokenType.WHITESPACE:
                idx += 1
            return idx

        def _read_group(idx: int, open_t: TokenType, close_t: TokenType) -> tuple[list[Token] | None, int]:
            j = _skip_ws(idx)
            if j >= len(tokens) or tokens[j].type != open_t:
                return None, idx
            j += 1
            depth = 1
            inner: list[Token] = []
            while j < len(tokens) and depth > 0:
                t = tokens[j]
                if t.type == open_t:
                    depth += 1
                    inner.append(t)
                elif t.type == close_t:
                    depth -= 1
                    if depth > 0:
                        inner.append(t)
                else:
                    inner.append(t)
                j += 1
            return inner, j

        parts: list[str] = []
        i = 0
        while i < len(tokens):
            tok = tokens[i]
            if tok.type == TokenType.TEXT:
                parts.append(tok.value)
            elif tok.type == TokenType.WHITESPACE:
                parts.append(" ")
            elif tok.type == TokenType.MATH_INLINE:
                parts.append(_latex_math_to_text(tok.extra.get("content", tok.value)))
            elif tok.type == TokenType.MATH_DISPLAY:
                parts.append(_latex_math_to_text(tok.extra.get("content", tok.value)))
            elif tok.type == TokenType.COMMAND:
                name = tok.extra.get("name", "")

                if name in _SYMBOL_MAP:
                    parts.append(_SYMBOL_MAP[name])
                    i += 1
                    continue

                if name in ("textbf", "textit", "emph", "underline", "text", "textrm", "texttt"):
                    inner, j = _read_group(i + 1, TokenType.BRACE_OPEN, TokenType.BRACE_CLOSE)
                    if inner is not None:
                        parts.append(self._tokens_to_text(inner))
                        i = j
                        continue

                if name in ("ref", "autoref", "cref", "Cref", "pageref", "eqref"):
                    inner, j = _read_group(i + 1, TokenType.BRACE_OPEN, TokenType.BRACE_CLOSE)
                    if inner is not None:
                        key = self._tokens_to_text(inner).strip()
                        if name == "pageref" and self.tex_structure:
                            label = self.tex_structure.labels.get(key)
                            display = str(label.page) if label else key
                        elif self.tex_structure:
                            display = self.tex_structure.resolve_ref(key) or key
                        else:
                            display = key
                        if name == "eqref":
                            parts.append(f"({display})")
                        else:
                            parts.append(display)
                        i = j
                        continue

                if name in ("cite", "citep", "citet", "citealt", "citealp", "citenum", "parencite", "textcite"):
                    j = i + 1
                    while True:
                        _, next_j = _read_group(j, TokenType.BRACKET_OPEN, TokenType.BRACKET_CLOSE)
                        if next_j == j:
                            break
                        j = next_j
                    inner, j2 = _read_group(j, TokenType.BRACE_OPEN, TokenType.BRACE_CLOSE)
                    if inner is not None:
                        keys_str = self._tokens_to_text(inner).strip()
                        keys = [k.strip() for k in keys_str.split(",") if k.strip()]
                        if self.tex_structure:
                            resolved = self.tex_structure.resolve_citation_keys(keys)
                        else:
                            resolved = keys
                        joined = ", ".join(resolved)
                        if name in ("citenum", "citealt", "citealp"):
                            parts.append(joined)
                        else:
                            parts.append(f"[{joined}]")
                        i = j2
                        continue

                if name in ("heiti", "songti", "kaiti", "fangsong"):
                    i += 1
                    continue
                if name in ("quad", "qquad", "enspace", "thinspace"):
                    parts.append("  " if name == "qquad" else " ")
                    i += 1
                    continue
            i += 1

        return "".join(parts).strip()

    # ── Main conversion ──────────────────────────────────────────────

    def convert(self, latex_content: str) -> Document:
        """Convert LaTeX content to a python-docx Document."""
        self.tokens = list(tokenize(latex_content))
        self.pos = 0

        while self._peek().type != TokenType.EOF:
            self._process_token()

        # Ensure body's <w:sectPr> is the very last child — python-docx
        # (and _make_logo_paragraph) relies on this invariant.
        self._fix_body_sectpr_position()

        # Apply global fixes
        fix_styles(self.doc, self.profile)
        fix_page_layout(self.doc, self.metadata)
        add_page_numbers(self.doc)
        enable_update_fields(self.doc)

        return self.doc

    @property
    def footnotes(self) -> list[tuple[int, str]]:
        """Collected footnotes as ``(id, text)`` pairs."""
        return list(self._footnotes)

    def _fix_body_sectpr_position(self):
        """Move the body-level <w:sectPr> to the very end of the body.

        python-docx assumes the final <w:sectPr> is the last child of
        <w:body>.  If body.append() was used anywhere, elements may have
        been placed after the sectPr, breaking doc.add_paragraph() and
        doc.add_picture().
        """
        body = self.doc.element.body
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is not None:
            body.remove(sect_pr)
            body.append(sect_pr)

    def _process_token(self):
        """Process the next token."""
        tok = self._peek()

        if tok.type == TokenType.EOF:
            return

        if tok.type == TokenType.COMMENT:
            self._advance()
            return

        # Skip everything before \begin{document} (preamble)
        if not self._in_document:
            if tok.type == TokenType.ENV_BEGIN and tok.extra.get("name") == "document":
                self._handle_env_begin()
            else:
                self._advance()
            return

        if tok.type == TokenType.ENV_BEGIN:
            self._handle_env_begin()
            return

        if tok.type == TokenType.ENV_END:
            self._handle_env_end()
            return

        if tok.type == TokenType.COMMAND:
            self._handle_command()
            return

        if tok.type == TokenType.PAR_BREAK:
            self._advance()
            self._finish_paragraph()
            return

        if tok.type == TokenType.WHITESPACE:
            self._advance()
            if self._current_para is not None:
                self._add_run(" ")
            return

        if tok.type == TokenType.TEXT:
            self._advance()
            self._add_run(tok.value)
            return

        if tok.type in (TokenType.MATH_INLINE, TokenType.MATH_DISPLAY):
            self._handle_math(tok)
            return

        if tok.type == TokenType.BRACE_OPEN:
            # Bare group — push format state
            self._advance()
            self.format_stack.append(self.format_stack[-1].copy())
            return

        if tok.type == TokenType.BRACE_CLOSE:
            self._advance()
            if len(self.format_stack) > 1:
                self.format_stack.pop()
            return

        if tok.type == TokenType.NEWLINE_CMD:
            self._advance()
            # In table context this is handled by the table builder
            if not self._env_stack or self._env_stack[-1] not in ("tabular", "tabular*", "tabularx", "longtable"):
                self._finish_paragraph()
            return

        if tok.type in (TokenType.AMPERSAND, TokenType.TOPRULE, TokenType.MIDRULE,
                        TokenType.BOTTOMRULE, TokenType.HLINE, TokenType.CMIDRULE):
            # Table tokens outside a table env — skip
            self._advance()
            return

        if tok.type == TokenType.ITEM:
            self._handle_item()
            return

        # Catch-all: skip
        self._advance()

    # ── Environment handlers ─────────────────────────────────────────

    def _handle_env_begin(self):
        tok = self._advance()
        env_name = tok.extra.get("name", "")

        if env_name == "document":
            self._in_document = True
            self._env_stack.append(env_name)
            return

        self._env_stack.append(env_name)

        # Consume optional placement arg like [ht], [htbp], etc.
        if env_name in ("table", "table*", "figure", "figure*"):
            self._read_optional_arg()

        if env_name in ("itemize", "enumerate"):
            self._list_depth += 1
            self._list_type_stack.append(env_name)
            self._list_counters.append(0)
            return

        if env_name in ("table", "table*"):
            self._finish_paragraph()
            self._table_count += 1
            self._pending_caption = None
            self._pending_label = None
            return

        if env_name in ("tabular", "tabular*", "tabularx", "longtable"):
            self._handle_tabular_env()
            return

        if env_name in ("figure", "figure*"):
            self._finish_paragraph()
            self._figure_count += 1
            self._pending_caption = None
            self._pending_label = None
            return

        if env_name in ("equation", "equation*", "align", "align*",
                         "gather", "gather*", "multline", "multline*"):
            self._handle_math_env(env_name)
            return

        if env_name == "abstract":
            self._finish_paragraph()
            self._add_heading_no_toc(self.profile.labels.abstract, 1)
            self._finish_paragraph()
            return

        if env_name in ("center", "flushleft", "flushright"):
            alignment_map = {
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "flushleft": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "flushright": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            }
            self._finish_paragraph()
            self._para_alignment = alignment_map.get(env_name)
            return

        if env_name in ("quote", "quotation"):
            self._finish_paragraph()
            return

        if env_name in ("verbatim", "lstlisting", "minted"):
            self._handle_verbatim_env(env_name)
            return

        # Unknown environment — just continue processing content inside

    def _handle_env_end(self):
        tok = self._advance()
        env_name = tok.extra.get("name", "")

        if env_name == "document":
            self._in_document = False
            if self._env_stack and self._env_stack[-1] == "document":
                self._env_stack.pop()
            return

        if env_name in ("itemize", "enumerate"):
            self._list_depth = max(0, self._list_depth - 1)
            if self._list_type_stack:
                self._list_type_stack.pop()
            if self._list_counters:
                self._list_counters.pop()
            self._finish_paragraph()
            if self._env_stack and self._env_stack[-1] == env_name:
                self._env_stack.pop()
            return

        if env_name in ("table", "table*"):
            self._finish_paragraph()
            self._after_float = True
            if self._env_stack and self._env_stack[-1] == env_name:
                self._env_stack.pop()
            return

        if env_name in ("figure", "figure*"):
            self._finish_paragraph()
            self._after_float = True
            if self._env_stack and self._env_stack[-1] == env_name:
                self._env_stack.pop()
            return

        if env_name in ("center", "flushleft", "flushright"):
            self._finish_paragraph()
            self._para_alignment = None
            if self._env_stack and self._env_stack[-1] == env_name:
                self._env_stack.pop()
            return

        if env_name in ("quote", "quotation"):
            self._finish_paragraph()
            if self._env_stack and self._env_stack[-1] == env_name:
                self._env_stack.pop()
            return

        if env_name == "abstract":
            self._finish_paragraph()
            if self._env_stack and self._env_stack[-1] == env_name:
                self._env_stack.pop()
            return

        if self._env_stack and self._env_stack[-1] == env_name:
            self._env_stack.pop()

        self._finish_paragraph()

    # ── Command handlers ─────────────────────────────────────────────

    def _handle_command(self):
        tok = self._advance()
        name = tok.extra.get("name", "")

        # Headings
        if name in ("chapter", "chapter*"):
            starred = name.endswith("*")
            self._add_heading(1, starred=starred)
            return
        if name in ("section", "section*"):
            starred = name.endswith("*")
            level = 2 if self.doc_class_type in ("report", "book") else 1
            self._add_heading(level, starred=starred)
            return
        if name in ("subsection", "subsection*"):
            starred = name.endswith("*")
            level = 3 if self.doc_class_type in ("report", "book") else 2
            self._add_heading(level, starred=starred)
            return
        if name in ("subsubsection", "subsubsection*"):
            starred = name.endswith("*")
            level = 4 if self.doc_class_type in ("report", "book") else 3
            self._add_heading(level, starred=starred)
            return

        # Text formatting
        if name == "textbf":
            self._handle_inline_format(bold=True)
            return
        if name in ("textit", "emph"):
            self._handle_inline_format(italic=True)
            return
        if name == "underline":
            self._handle_inline_format(underline=True)
            return
        if name == "textsuperscript":
            self._handle_inline_format(superscript=True)
            return
        if name == "textsubscript":
            self._handle_inline_format(subscript=True)
            return
        if name == "texttt":
            self._handle_inline_format(font_name=self.profile.fonts.monospace)
            return
        if name == "textrm":
            self._handle_inline_format()
            return
        if name == "text":
            self._handle_inline_format()
            return

        # Symbol commands → Unicode
        if name in _SYMBOL_MAP:
            self._add_run(_SYMBOL_MAP[name])
            return

        # Spacing commands that produce whitespace
        if name in ("quad", "qquad", "enspace", "thinspace"):
            self._add_run("  " if name == "qquad" else " ")
            return
        if name == ",":
            self._add_run(" ")
            return

        # CJK font commands — data-driven lookup
        cjk_font = self.profile.get_cjk_font(name)
        if cjk_font is not None:
            self._handle_font_switch(cjk_font)
            return

        # Page breaks
        if name in ("newpage", "clearpage", "cleardoublepage"):
            self._finish_paragraph()
            self.doc.add_page_break()
            self._finish_paragraph()
            return

        # Caption
        if name == "caption":
            self._handle_caption()
            return

        # Label — skip
        if name == "label":
            self._read_brace_group()
            return

        # Cross-references
        if name == "ref":
            ref_key = self._read_brace_group()
            if self.tex_structure:
                resolved = self.tex_structure.resolve_ref(ref_key)
                self._add_run(resolved if resolved else f"[{ref_key}]")
            else:
                self._add_run(f"[{ref_key}]")
            return
        if name == "eqref":
            ref_key = self._read_brace_group()
            if self.tex_structure:
                resolved = self.tex_structure.resolve_ref(ref_key)
                self._add_run(f"({resolved})" if resolved else f"({ref_key})")
            else:
                self._add_run(f"({ref_key})")
            return
        if name == "pageref":
            ref_key = self._read_brace_group()
            if self.tex_structure:
                label = self.tex_structure.labels.get(ref_key)
                self._add_run(str(label.page) if label else f"[{ref_key}]")
            else:
                self._add_run(f"[{ref_key}]")
            return
        if name in ("cite", "citep", "citet", "citealt", "citealp", "citenum", "parencite", "textcite"):
            # Consume one or more optional args (e.g. biblatex prenote/postnote).
            while self._read_optional_arg() is not None:
                pass
            cite_keys_str = self._read_brace_group()
            if self.tex_structure:
                keys = [k.strip() for k in cite_keys_str.split(",") if k.strip()]
                resolved = self.tex_structure.resolve_citation_keys(keys)
                joined = ", ".join(resolved)
                if name in ("citenum", "citealt", "citealp"):
                    self._add_run(joined)
                else:
                    self._add_run(f"[{joined}]")
            else:
                if name in ("citenum", "citealt", "citealp"):
                    self._add_run(cite_keys_str)
                else:
                    self._add_run(f"[{cite_keys_str}]")
            return
        if name in ("autoref", "cref", "Cref"):
            ref_key = self._read_brace_group()
            if self.tex_structure:
                resolved = self.tex_structure.resolve_ref(ref_key)
                self._add_run(resolved if resolved else ref_key)
            else:
                self._add_run(ref_key)
            return
        if name == "bibliography":
            self._read_brace_group()
            return

        # Template keyword commands
        if name in ("keywords", "KEYWORDS"):
            kw = self._read_brace_group().strip()
            if name == "keywords":
                prefix = self.profile.labels.keywords_zh_prefix
            else:
                prefix = self.profile.labels.keywords_en_prefix
            bold_fmt = self.format_stack[-1].merge(bold=True)
            self._add_run(prefix, bold_fmt)
            if kw:
                self._add_run(kw)
            return

        # Footnote
        if name == "footnote":
            self._handle_footnote()
            return

        # URL/href
        if name == "url":
            url_text = self._read_brace_group()
            self._add_hyperlink(url_text, url_text)
            return
        if name == "href":
            url = self._read_brace_group()
            display = self._read_brace_group()
            self._add_hyperlink(url, display)
            return

        # Images
        if name == "includegraphics":
            self._handle_includegraphics()
            return

        # Table of contents
        if name == "tableofcontents":
            self._handle_toc()
            return

        # List of figures / List of tables
        if name == "listoffigures":
            self._handle_list_of("figure")
            return
        if name == "listoftables":
            self._handle_list_of("table")
            return

        # Vertical spacing
        if name == "vspace":
            self._read_optional_arg()  # *
            self._read_brace_group()
            return
        if name == "hspace":
            self._read_optional_arg()
            self._read_brace_group()
            self._add_run(" ")
            return

        # Font size commands with argument: \fontsize{size}{skip}
        if name == "fontsize":
            self._read_brace_group()
            self._read_brace_group()
            return

        # Skip known no-arg commands
        if name in self._SKIP_COMMANDS:
            return

        # Skip known commands with arguments
        if name in self._SKIP_WITH_ARG:
            # Consume optional arg then brace group(s)
            self._read_optional_arg()
            self._read_brace_group()
            # Some commands take multiple args
            if name in ("setlength", "addtolength", "renewcommand",
                        "newcommand", "providecommand", "setcounter",
                        "addtocounter", "DeclareCaptionFont"):
                self._read_brace_group()
            if name in ("newcommand", "providecommand"):
                self._read_optional_arg()
                self._read_brace_group()
            return

        # usepackage — skip
        if name == "usepackage":
            self._read_optional_arg()
            self._read_brace_group()
            return

        # documentclass — skip
        if name == "documentclass":
            self._read_optional_arg()
            self._read_brace_group()
            return

        # input/include — skip
        if name in ("input", "include"):
            self._read_brace_group()
            return

        # multicolumn — used in tables, handle inline
        if name == "multicolumn":
            self._handle_multicolumn()
            return

        # Unknown command: try to consume its argument and output text
        self._skip_whitespace()
        if self._peek().type == TokenType.BRACE_OPEN:
            # Has argument — consume it and process as inline text
            inner_tokens = self._read_brace_group_tokens()
            self._process_inline_tokens(inner_tokens)
        # else: no argument, just skip the command

    # ── Inline format handling ───────────────────────────────────────

    def _handle_inline_format(self, **kwargs):
        """Handle \textbf{...}, \textit{...} etc."""
        new_fmt = self.format_stack[-1].merge(**kwargs)
        self.format_stack.append(new_fmt)
        inner_tokens = self._read_brace_group_tokens()
        self._process_inline_tokens(inner_tokens)
        self.format_stack.pop()

    def _handle_font_switch(self, font_name: str):
        r"""Handle \\heiti, \\songti etc. as font switch.

        If followed by a brace group, apply to that group.
        Otherwise, modify current format state.
        """
        self._skip_whitespace()
        if self._peek().type == TokenType.BRACE_OPEN:
            new_fmt = self.format_stack[-1].merge(font_name=font_name)
            self.format_stack.append(new_fmt)
            inner_tokens = self._read_brace_group_tokens()
            self._process_inline_tokens(inner_tokens)
            self.format_stack.pop()
        else:
            # Font switch affects the rest of the current group
            self.format_stack[-1].font_name = font_name

    def _process_inline_tokens(self, tokens: list[Token]):
        """Process tokens inline (within the current paragraph)."""
        saved_tokens = self.tokens
        saved_pos = self.pos
        self.tokens = tokens
        self.pos = 0

        while self._peek().type != TokenType.EOF:
            tok = self._peek()
            if tok.type == TokenType.TEXT:
                self._advance()
                self._add_run(tok.value)
            elif tok.type == TokenType.WHITESPACE:
                self._advance()
                self._add_run(" ")
            elif tok.type == TokenType.COMMAND:
                self._handle_command()
            elif tok.type == TokenType.BRACE_OPEN:
                self._advance()
                self.format_stack.append(self.format_stack[-1].copy())
            elif tok.type == TokenType.BRACE_CLOSE:
                self._advance()
                if len(self.format_stack) > 1:
                    self.format_stack.pop()
            elif tok.type in (TokenType.MATH_INLINE, TokenType.MATH_DISPLAY):
                self._handle_math(tok)
            else:
                self._advance()

        self.tokens = saved_tokens
        self.pos = saved_pos

    # ── List items ───────────────────────────────────────────────────

    def _handle_item(self):
        tok = self._advance()
        self._finish_paragraph()

        label = tok.extra.get("label")
        list_type = self._list_type_stack[-1] if self._list_type_stack else "itemize"

        if list_type == "enumerate" and self._list_counters:
            self._list_counters[-1] += 1
            num = self._list_counters[-1]
            prefix = label if label else f"{num}. "
        else:
            prefix = label if label else "\u25cf "  # ● filled circle (matches LaTeX)

        para = self._ensure_paragraph()
        # Hanging indent is already set by _apply_list_indent in _ensure_paragraph
        # Add prefix
        if prefix:
            run = para.add_run(prefix if prefix.endswith(" ") else f"{prefix} ")

    # ── Caption handling ─────────────────────────────────────────────

    def _add_seq_field(self, para, identifier: str, display_value: str,
                       bold: bool = False, font_size_pt: float = 10.5):
        """Insert a Word SEQ field into *para*.

        The ``SEQ`` field makes the caption discoverable by Word's
        ``TOC \\c`` field, so that List of Figures / List of Tables
        can be populated automatically.
        """
        def _rpr():
            rPr = OxmlElement("w:rPr")
            if bold:
                rPr.append(OxmlElement("w:b"))
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(font_size_pt * 2)))
            rPr.append(sz)
            return rPr

        # fldChar begin
        r1 = OxmlElement("w:r")
        r1.append(_rpr())
        fc1 = OxmlElement("w:fldChar")
        fc1.set(qn("w:fldCharType"), "begin")
        r1.append(fc1)
        para._element.append(r1)

        # instrText
        r2 = OxmlElement("w:r")
        r2.append(_rpr())
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" SEQ {identifier} \\* ARABIC "
        r2.append(instr)
        para._element.append(r2)

        # fldChar separate
        r3 = OxmlElement("w:r")
        r3.append(_rpr())
        fc2 = OxmlElement("w:fldChar")
        fc2.set(qn("w:fldCharType"), "separate")
        r3.append(fc2)
        para._element.append(r3)

        # display value
        r4 = OxmlElement("w:r")
        r4.append(_rpr())
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = display_value
        r4.append(t)
        para._element.append(r4)

        # fldChar end
        r5 = OxmlElement("w:r")
        r5.append(_rpr())
        fc3 = OxmlElement("w:fldChar")
        fc3.set(qn("w:fldCharType"), "end")
        r5.append(fc3)
        para._element.append(r5)

    def _handle_caption(self):
        caption_text = _normalize_latex_text(self._read_brace_group())
        # Determine if we're in a figure or table
        in_figure = any(e in ("figure", "figure*") for e in self._env_stack)
        in_table = any(e in ("table", "table*") for e in self._env_stack)

        self._finish_paragraph()
        para = self._ensure_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.paragraph_format.first_line_indent = Pt(0)

        if in_figure:
            fig_num = str(self._figure_count)
            if self.tex_structure:
                entry = self.tex_structure.find_figure(self._figure_count)
                if entry:
                    fig_num = entry.number
            label = self.profile.labels.figure_prefix
            if self.tex_structure:
                # TeX number as plain text — SEQ field would auto-renumber
                run = para.add_run(f"{label} {fig_num}")
                run.bold = True
                run.font.size = Pt(10.5)
            else:
                run = para.add_run(f"{label} ")
                run.bold = True
                run.font.size = Pt(10.5)
                self._add_seq_field(para, label, fig_num, bold=True)
        elif in_table:
            tab_num = str(self._table_count)
            if self.tex_structure:
                entry = self.tex_structure.find_table(self._table_count)
                if entry:
                    tab_num = entry.number
            label = self.profile.labels.table_prefix
            if self.tex_structure:
                run = para.add_run(f"{label} {tab_num}")
                run.bold = True
                run.font.size = Pt(10.5)
            else:
                run = para.add_run(f"{label} ")
                run.bold = True
                run.font.size = Pt(10.5)
                self._add_seq_field(para, label, tab_num, bold=True)

        run = para.add_run(f"  {caption_text}")
        run.font.size = Pt(10.5)

        try:
            para.style = self.doc.styles["Caption"]
        except KeyError:
            pass

        # Force black on ALL runs — Caption style defaults to blue (accent1)
        for r in para.runs:
            r.font.color.rgb = RGBColor(0, 0, 0)

        self._finish_paragraph()

    # ── Footnote ─────────────────────────────────────────────────────

    def _handle_footnote(self):
        note_tokens = self._read_brace_group_tokens()
        note_text = self._tokens_to_text(note_tokens).strip()
        if not note_text:
            return

        self._footnote_count += 1
        footnote_id = self._footnote_count
        self._footnotes.append((footnote_id, note_text))

        para = self._ensure_paragraph()
        run = para.add_run()
        rPr = run._element.get_or_add_rPr()
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "FootnoteReference")
        rPr.append(rStyle)
        footnote_ref = OxmlElement("w:footnoteReference")
        footnote_ref.set(qn("w:id"), str(footnote_id))
        run._element.append(footnote_ref)

    # ── Image handling ───────────────────────────────────────────────

    def _handle_includegraphics(self):
        opts = self._read_optional_arg()
        img_path_str = self._read_brace_group()

        # Parse width from options
        width = None
        if opts:
            m = re.search(r"width\s*=\s*([\d.]+)\s*(cm|mm|in|\\textwidth|\\linewidth|\\columnwidth)", opts)
            if m:
                num = float(m.group(1))
                unit = m.group(2)
                if unit == "cm":
                    width = Cm(num)
                elif unit == "mm":
                    width = Cm(num / 10)
                elif unit == "in":
                    width = Cm(num * 2.54)
                elif unit in ("\\textwidth", "\\linewidth", "\\columnwidth"):
                    width = Cm(num * 15)  # approximate text width

        # Resolve image path
        img_path = self._resolve_image_path(img_path_str)
        if img_path is None:
            self._add_run(f"[Image: {img_path_str}]")
            return

        self._finish_paragraph()
        try:
            if width:
                self.doc.add_picture(str(img_path), width=width)
            else:
                self.doc.add_picture(str(img_path), width=Cm(12))
            # Center the image and remove first-line indent
            last_para = self.doc.paragraphs[-1]
            last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            last_para.paragraph_format.first_line_indent = Pt(0)
        except Exception as e:
            logger.warning("Failed to insert image %s: %s", img_path_str, e)
            para = self.doc.add_paragraph()
            para.add_run(f"[Image: {img_path_str}]")
        self._finish_paragraph()

    def _resolve_image_path(self, path_str: str) -> Path | None:
        """Resolve an image path relative to image_base_dir.

        Also searches subdirectories commonly used by LaTeX templates
        (e.g. ``Img/`` declared via ``\\graphicspath``).
        """
        extensions = ["", ".png", ".jpg", ".jpeg", ".pdf", ".eps", ".svg"]
        # Search dirs: base dir + immediate subdirectories (covers \graphicspath)
        search_dirs = [self.image_base_dir]
        try:
            search_dirs += [
                d for d in self.image_base_dir.iterdir() if d.is_dir()
            ]
        except OSError:
            pass
        for directory in search_dirs:
            for ext in extensions:
                candidate = directory / f"{path_str}{ext}"
                if candidate.exists():
                    if candidate.suffix.lower() == ".pdf":
                        return self._convert_pdf_image(candidate)
                    return candidate
        return None

    def _convert_pdf_image(self, pdf_path: Path) -> Path | None:
        """Convert PDF image to PNG for inclusion in DOCX."""
        try:
            import fitz
            import tempfile
            doc = fitz.open(str(pdf_path))
            page = doc[0]
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
            tmp = Path(tempfile.mktemp(suffix=".png"))
            pix.save(str(tmp))
            doc.close()
            return tmp
        except Exception as e:
            logger.warning("Failed to convert PDF image: %s", e)
            return None

    # ── Math handling ───────────────────────────────────────────────

    def _handle_math(self, tok: Token):
        from .math_handler import add_math_to_paragraph
        self._advance()
        content = tok.extra.get("content", tok.value)

        if tok.type == TokenType.MATH_DISPLAY:
            self._finish_paragraph()
            para = self._ensure_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            add_math_to_paragraph(para, content, display=True)
            self._finish_paragraph()
        else:
            # Inline math
            para = self._ensure_paragraph()
            add_math_to_paragraph(para, content, display=False)

    def _handle_math_env(self, env_name: str):
        """Handle math environments (equation, align, etc.)."""
        from .math_handler import add_math_to_paragraph

        numbered = not env_name.endswith("*")
        # align/gather: each \\ line gets its own number
        multi_line = env_name.rstrip("*") in ("align", "gather")

        # Collect all tokens until \end{env_name}
        parts = []
        while self._peek().type != TokenType.EOF:
            tok = self._peek()
            if tok.type == TokenType.ENV_END and tok.extra.get("name") == env_name:
                self._advance()
                if self._env_stack and self._env_stack[-1] == env_name:
                    self._env_stack.pop()
                break
            self._advance()
            parts.append(tok.value)

        raw = "".join(parts).strip()

        # Split multi-line environments by \\
        if multi_line:
            import re
            lines = [ln.strip() for ln in re.split(r"\\\\", raw) if ln.strip()]
            # Remove \nonumber / \notag markers and track which lines are numbered
            clean_lines = []
            for ln in lines:
                has_nonumber = ("\\nonumber" in ln or "\\notag" in ln)
                ln = ln.replace("\\nonumber", "").replace("\\notag", "").strip()
                # Remove alignment & markers for display
                ln = ln.replace("&", " ").strip()
                if ln:
                    clean_lines.append((ln, numbered and not has_nonumber))
        else:
            clean_lines = [(raw, numbered)]

        self._finish_paragraph()

        for math_content, should_number in clean_lines:
            para = self._ensure_paragraph()
            para.paragraph_format.first_line_indent = Pt(0)

            if should_number:
                self._equation_count += 1
                self._add_numbered_equation(para, math_content)
            else:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                add_math_to_paragraph(para, math_content, display=True)

            self._finish_paragraph()

    def _add_numbered_equation(self, para, math_content: str):
        """Add a display equation with right-aligned number: [math] ... (N)."""
        from .math_handler import add_math_to_paragraph

        pPr = para._element.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab_c = OxmlElement("w:tab")
        tab_c.set(qn("w:val"), "center")
        tab_c.set(qn("w:pos"), "4536")   # center of text area
        tabs.append(tab_c)
        tab_r = OxmlElement("w:tab")
        tab_r.set(qn("w:val"), "right")
        tab_r.set(qn("w:pos"), "9072")   # right margin
        tabs.append(tab_r)
        pPr.append(tabs)

        para.add_run("\t")
        add_math_to_paragraph(para, math_content, display=True)
        chap = self.counters.chapter
        eq_num = (f"({chap}.{self._equation_count})"
                  if self.doc_class_type in ("report", "book") and chap
                  else f"({self._equation_count})")
        run = para.add_run(f"\t{eq_num}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # ── Verbatim / code ──────────────────────────────────────────────

    def _handle_verbatim_env(self, env_name: str):
        """Handle verbatim-like environments."""
        parts = []
        while self._peek().type != TokenType.EOF:
            tok = self._peek()
            if tok.type == TokenType.ENV_END and tok.extra.get("name") == env_name:
                self._advance()
                if self._env_stack and self._env_stack[-1] == env_name:
                    self._env_stack.pop()
                break
            self._advance()
            parts.append(tok.value)

        content = "".join(parts).strip()
        self._finish_paragraph()
        for line in content.split("\n"):
            para = self.doc.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(0)
            run = para.add_run(line)
            run.font.name = self.profile.fonts.monospace
            run.font.size = Pt(10)
        self._finish_paragraph()

    # ── Table handling (basic — Phase 3 enhances this) ───────────────

    def _handle_tabular_env(self):
        """Handle tabular environment — build a Word table."""
        from .table_builder import build_table

        # Read column spec
        col_spec = self._read_brace_group()

        # Collect all tokens until \end{tabular}
        env_name = self._env_stack[-1] if self._env_stack else "tabular"
        table_tokens = []
        while self._peek().type != TokenType.EOF:
            tok = self._peek()
            if tok.type == TokenType.ENV_END and tok.extra.get("name") == env_name:
                self._advance()
                if self._env_stack and self._env_stack[-1] == env_name:
                    self._env_stack.pop()
                break
            table_tokens.append(self._advance())

        self._finish_paragraph()
        build_table(self.doc, col_spec, table_tokens, self)
        self._finish_paragraph()

    # ── TOC ──────────────────────────────────────────────────────────

    def _handle_toc(self):
        """Insert a Table of Contents.

        If TeX ``.aux`` data is available the TOC entries are rendered
        directly with page numbers (matching the PDF).  Otherwise a Word
        ``TOC`` field placeholder is inserted.
        """
        self._finish_paragraph()
        self._add_heading_no_toc(self.profile.labels.toc, 1)
        self._finish_paragraph()

        # Render from .aux data if available
        if self.tex_structure and self.tex_structure.toc_entries:
            level_indent = {"chapter": 0, "section": 0, "subsection": 1, "subsubsection": 2}
            level_size = {"chapter": 12, "section": 12, "subsection": 11, "subsubsection": 10.5}
            for entry in self.tex_structure.toc_entries:
                para = self.doc.add_paragraph()
                para.paragraph_format.first_line_indent = Pt(0)
                indent_level = level_indent.get(entry.level, 0)
                para.paragraph_format.left_indent = Cm(indent_level * 0.75)
                font_size = level_size.get(entry.level, 11)

                # Right-aligned tab stop with dot leader for page number
                pPr = para._element.get_or_add_pPr()
                tabs = OxmlElement("w:tabs")
                tab = OxmlElement("w:tab")
                tab.set(qn("w:val"), "right")
                tab.set(qn("w:leader"), "dot")
                tab.set(qn("w:pos"), "9072")
                tabs.append(tab)
                pPr.append(tabs)

                # Title text
                run = para.add_run(entry.full_title)
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if entry.level in ("chapter",):
                    run.bold = True
                # Tab + page number
                run = para.add_run(f"\t{entry.page}")
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0, 0, 0)
            self._finish_paragraph()
            return

        # Fallback: Word TOC field
        toc_para = make_toc_field_paragraph(
            hint_text=self.profile.labels.toc_update_hint,
        )
        self._append_element(toc_para)
        self._finish_paragraph()

    def _handle_list_of(self, kind: str):
        """Insert a List of Figures or List of Tables.

        *kind*: ``"figure"`` or ``"table"``.

        If TeX ``.aux`` data is available the entries are rendered directly
        as static paragraphs (matching the PDF).  Otherwise a Word ``TOC``
        field is inserted that the user can update after opening the file.
        """
        if kind == "figure":
            heading_text = self.profile.labels.list_of_figures
            label = self.profile.labels.figure_prefix
        else:
            heading_text = self.profile.labels.list_of_tables
            label = self.profile.labels.table_prefix

        self._finish_paragraph()
        self._add_heading_no_toc(heading_text, 1)
        self._finish_paragraph()

        # If .aux data is available, render entries directly
        if self.tex_structure:
            entries = (self.tex_structure.lof_entries if kind == "figure"
                       else self.tex_structure.lot_entries)
            if entries:
                for entry in entries:
                    para = self.doc.add_paragraph()
                    para.paragraph_format.first_line_indent = Pt(0)
                    # Right-aligned tab stop with dot leader for page number
                    pPr = para._element.get_or_add_pPr()
                    tabs = OxmlElement("w:tabs")
                    tab = OxmlElement("w:tab")
                    tab.set(qn("w:val"), "right")
                    tab.set(qn("w:leader"), "dot")
                    tab.set(qn("w:pos"), "9072")  # ~16cm in twips
                    tabs.append(tab)
                    pPr.append(tabs)

                    run = para.add_run(f"{label} {entry.number}")
                    run.bold = True
                    run.font.size = Pt(12)
                    run = para.add_run(f"  {entry.caption}")
                    run.font.size = Pt(12)
                    # Tab + page number
                    run = para.add_run(f"\t{entry.page}")
                    run.font.size = Pt(12)
                self._finish_paragraph()
                return

        # Fallback: Word TOC field (requires field update in Word)
        from app.core.compiler.word_postprocessor import _make_list_field_paragraph
        lof_para = _make_list_field_paragraph(kind, label=label)
        self._append_element(lof_para)
        self._finish_paragraph()

    def _append_element(self, element):
        """Append an OxmlElement to the body before the final <w:sectPr>.

        Using body.append() directly would place the element after the
        body-level sectPr, breaking python-docx's internal assumptions.
        """
        body = self.doc.element.body
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is not None:
            sect_pr.addprevious(element)
        else:
            body.append(element)

    # ── Multicolumn ──────────────────────────────────────────────────

    def _handle_multicolumn(self):
        r"""Handle \\multicolumn{n}{spec}{content} — just output content as text."""
        self._read_brace_group()  # n
        self._read_brace_group()  # column spec
        content = self._read_brace_group()
        self._add_run(content)
