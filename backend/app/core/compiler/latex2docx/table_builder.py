"""Table builder for LaTeX → DOCX conversion.

Parses tabular column specs, handles booktabs/hline borders, and builds
python-docx tables with correct column widths and border styles.
"""

import re
import logging
from dataclasses import dataclass, field
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Emu, RGBColor

from .tokenizer import Token, TokenType
from .text_utils import normalize_latex_text, SYMBOL_MAP

if TYPE_CHECKING:
    from .converter import LatexToDocxConverter

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Column definition
# ---------------------------------------------------------------------------

@dataclass
class ColumnDef:
    """Definition of a single table column."""
    align: str = "left"        # "left", "center", "right"
    width_cm: float | None = None  # fixed width in cm, None = auto
    left_border: bool = False
    right_border: bool = False


def parse_column_spec(spec: str) -> list[ColumnDef]:
    """Parse a LaTeX column specification like ``|l|c|r|p{4cm}|``.

    Returns a list of ColumnDef objects.
    """
    columns = []
    i = 0
    pending_left_border = False

    while i < len(spec):
        ch = spec[i]

        if ch == "|":
            if columns:
                columns[-1].right_border = True
            else:
                pending_left_border = True
            i += 1
            continue

        if ch in ("l", "c", "r"):
            align_map = {"l": "left", "c": "center", "r": "right"}
            col = ColumnDef(align=align_map[ch])
            if pending_left_border:
                col.left_border = True
                pending_left_border = False
            columns.append(col)
            i += 1
            continue

        if ch in ("p", "m", "b"):
            # p{width}, m{width}, b{width}
            i += 1
            # Skip to {
            while i < len(spec) and spec[i] != "{":
                i += 1
            if i < len(spec):
                i += 1  # skip {
                width_str = ""
                depth = 1
                while i < len(spec) and depth > 0:
                    if spec[i] == "{":
                        depth += 1
                    elif spec[i] == "}":
                        depth -= 1
                        if depth == 0:
                            i += 1
                            break
                    width_str += spec[i]
                    i += 1
                width_cm = _parse_width(width_str)
                col = ColumnDef(align="left", width_cm=width_cm)
                if pending_left_border:
                    col.left_border = True
                    pending_left_border = False
                columns.append(col)
            continue

        if ch == "X":
            # tabularx X column — auto width, left aligned
            col = ColumnDef(align="left")
            if pending_left_border:
                col.left_border = True
                pending_left_border = False
            columns.append(col)
            i += 1
            continue

        if ch == "@" or ch == "!":
            # @{...} or !{...} — inter-column material, skip
            i += 1
            if i < len(spec) and spec[i] == "{":
                depth = 1
                i += 1
                while i < len(spec) and depth > 0:
                    if spec[i] == "{":
                        depth += 1
                    elif spec[i] == "}":
                        depth -= 1
                    i += 1
            continue

        if ch == "*":
            # *{n}{spec} — repeated columns
            i += 1
            # Read count
            if i < len(spec) and spec[i] == "{":
                i += 1
                count_str = ""
                while i < len(spec) and spec[i] != "}":
                    count_str += spec[i]
                    i += 1
                i += 1  # skip }
                try:
                    count = int(count_str)
                except ValueError:
                    count = 1
                # Read sub-spec
                if i < len(spec) and spec[i] == "{":
                    i += 1
                    sub_spec = ""
                    depth = 1
                    while i < len(spec) and depth > 0:
                        if spec[i] == "{":
                            depth += 1
                        elif spec[i] == "}":
                            depth -= 1
                            if depth == 0:
                                i += 1
                                break
                        sub_spec += spec[i]
                        i += 1
                    for _ in range(count):
                        sub_cols = parse_column_spec(sub_spec)
                        columns.extend(sub_cols)
            continue

        # Skip whitespace and other characters
        i += 1

    return columns


def _parse_width(width_str: str) -> float | None:
    """Parse a LaTeX width like '4cm', '2.5in', '0.3\\textwidth'."""
    width_str = width_str.strip()

    m = re.match(r"([\d.]+)\s*cm", width_str)
    if m:
        return float(m.group(1))

    m = re.match(r"([\d.]+)\s*mm", width_str)
    if m:
        return float(m.group(1)) / 10

    m = re.match(r"([\d.]+)\s*in", width_str)
    if m:
        return float(m.group(1)) * 2.54

    m = re.match(r"([\d.]+)\s*(?:pt|bp)", width_str)
    if m:
        return float(m.group(1)) / 72 * 2.54

    m = re.match(r"([\d.]+)\s*\\(?:textwidth|linewidth|columnwidth)", width_str)
    if m:
        return float(m.group(1)) * 15  # approximate page width

    return None


# ---------------------------------------------------------------------------
# Border style detection
# ---------------------------------------------------------------------------

@dataclass
class BorderStyle:
    """Describes the border style for a table."""
    style: str = "none"   # "three_line", "grid", "hline_only", "none"


def detect_border_style(tokens: list[Token], col_spec: str = "") -> BorderStyle:
    """Detect border style from table tokens and column spec.

    ``col_spec`` is the raw LaTeX column specification (e.g. ``|l|c|r|``
    or ``p{4cm}p{10cm}``).  When ``\\hline`` is used *without* ``|``
    separators the table gets horizontal-only borders (matching the PDF
    rendering).  With ``|`` separators it becomes a full grid.
    """
    has_booktabs = any(
        t.type in (TokenType.TOPRULE, TokenType.MIDRULE, TokenType.BOTTOMRULE)
        for t in tokens
    )
    has_hline = any(t.type == TokenType.HLINE for t in tokens)
    has_vert = "|" in col_spec

    if has_booktabs:
        return BorderStyle(style="three_line")
    elif has_hline:
        if has_vert:
            return BorderStyle(style="grid")
        else:
            return BorderStyle(style="hline_only")
    else:
        return BorderStyle(style="none")


# ---------------------------------------------------------------------------
# Parse table rows from tokens
# ---------------------------------------------------------------------------

@dataclass
class CellData:
    """Data for a single table cell."""
    tokens: list[Token] = field(default_factory=list)
    colspan: int = 1
    align: str | None = None  # override from \multicolumn


def parse_table_rows(tokens: list[Token]) -> list[list[CellData]]:
    """Parse table tokens into rows of cells."""
    rows = []
    current_row: list[CellData] = [CellData()]

    i = 0
    while i < len(tokens):
        tok = tokens[i]

        if tok.type in (TokenType.TOPRULE, TokenType.MIDRULE,
                        TokenType.BOTTOMRULE, TokenType.HLINE,
                        TokenType.CMIDRULE):
            i += 1
            continue

        if tok.type == TokenType.AMPERSAND:
            current_row.append(CellData())
            i += 1
            continue

        if tok.type == TokenType.NEWLINE_CMD:
            # End of row
            if any(c.tokens for c in current_row):
                rows.append(current_row)
            current_row = [CellData()]
            i += 1
            continue

        if tok.type == TokenType.COMMAND and tok.extra.get("name") == "multicolumn":
            # \multicolumn{n}{spec}{content}
            i += 1
            # Read n
            n_str = _read_brace_from_tokens(tokens, i)
            i = n_str[1]
            try:
                colspan = int(n_str[0])
            except ValueError:
                colspan = 1
            # Read spec
            spec = _read_brace_from_tokens(tokens, i)
            i = spec[1]
            align_map = {"l": "left", "c": "center", "r": "right"}
            align = None
            for ch in spec[0]:
                if ch in align_map:
                    align = align_map[ch]
                    break
            # Read content
            content = _read_brace_tokens_from_tokens(tokens, i)
            i = content[1]
            cell = CellData(tokens=content[0], colspan=colspan, align=align)
            current_row[-1] = cell
            continue

        # Regular token — add to current cell
        current_row[-1].tokens.append(tok)
        i += 1

    # Don't forget last row
    if any(c.tokens for c in current_row):
        rows.append(current_row)

    # Filter out rows that are entirely whitespace
    rows = [
        row for row in rows
        if any(
            any(t.type not in (TokenType.WHITESPACE,) for t in c.tokens)
            for c in row
        )
    ]

    return rows


def _read_brace_from_tokens(tokens: list[Token], pos: int) -> tuple[str, int]:
    """Read a {}-group from token list, return (text, new_pos)."""
    # Skip whitespace
    while pos < len(tokens) and tokens[pos].type == TokenType.WHITESPACE:
        pos += 1
    if pos >= len(tokens) or tokens[pos].type != TokenType.BRACE_OPEN:
        return ("", pos)
    pos += 1
    depth = 1
    parts = []
    while pos < len(tokens) and depth > 0:
        if tokens[pos].type == TokenType.BRACE_OPEN:
            depth += 1
            parts.append("{")
        elif tokens[pos].type == TokenType.BRACE_CLOSE:
            depth -= 1
            if depth > 0:
                parts.append("}")
        else:
            parts.append(tokens[pos].value)
        pos += 1
    return ("".join(parts), pos)


def _read_brace_tokens_from_tokens(tokens: list[Token], pos: int) -> tuple[list[Token], int]:
    """Read a {}-group from token list, return (token_list, new_pos)."""
    while pos < len(tokens) and tokens[pos].type == TokenType.WHITESPACE:
        pos += 1
    if pos >= len(tokens) or tokens[pos].type != TokenType.BRACE_OPEN:
        return ([], pos)
    pos += 1
    depth = 1
    result = []
    while pos < len(tokens) and depth > 0:
        if tokens[pos].type == TokenType.BRACE_OPEN:
            depth += 1
            result.append(tokens[pos])
        elif tokens[pos].type == TokenType.BRACE_CLOSE:
            depth -= 1
            if depth > 0:
                result.append(tokens[pos])
        else:
            result.append(tokens[pos])
        pos += 1
    return (result, pos)


# ---------------------------------------------------------------------------
# Build Word table
# ---------------------------------------------------------------------------

def build_table(
    doc: Document,
    col_spec: str,
    tokens: list[Token],
    converter: "LatexToDocxConverter",
) -> None:
    """Build a Word table from column spec and table tokens."""
    columns = parse_column_spec(col_spec)
    if not columns:
        # Fallback: estimate from first row
        columns = [ColumnDef(align="left")]

    border_style = detect_border_style(tokens, col_spec)
    rows_data = parse_table_rows(tokens)

    if not rows_data:
        return

    # Determine actual number of columns
    max_cols = max(
        sum(cell.colspan for cell in row)
        for row in rows_data
    )
    while len(columns) < max_cols:
        columns.append(ColumnDef(align="left"))

    num_rows = len(rows_data)
    num_cols = len(columns)

    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Force "Normal Table" (TableNormal) style — the plain base style with
    # no conditional formatting.  Without an explicit style reference, Word
    # for Mac auto-applies the last-used table style, which often has blue
    # first-row formatting.  We also disable all conditional format bands.
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is not None:
            tblStyle.set(qn("w:val"), "TableNormal")
        else:
            tblStyle = OxmlElement("w:tblStyle")
            tblStyle.set(qn("w:val"), "TableNormal")
            tblPr.insert(0, tblStyle)
        tblLook = tblPr.find(qn("w:tblLook"))
        if tblLook is not None:
            tblLook.set(qn("w:firstRow"), "0")
            tblLook.set(qn("w:lastRow"), "0")
            tblLook.set(qn("w:firstColumn"), "0")
            tblLook.set(qn("w:noHBand"), "1")
            tblLook.set(qn("w:val"), "0000")

    # Reduce cell margins so text sits closer to border lines,
    # matching the tight spacing of LaTeX tables.
    _apply_cell_margins(table)

    # Apply column widths
    _apply_column_widths(table, columns)

    # Apply borders
    _apply_borders(table, border_style, num_rows)

    # Fill cell content
    for ri, row_data in enumerate(rows_data):
        row = table.rows[ri]
        ci = 0
        for cell_data in row_data:
            if ci >= num_cols:
                break

            cell = row.cells[ci]

            # Handle multicolumn merge
            if cell_data.colspan > 1:
                end_ci = min(ci + cell_data.colspan - 1, num_cols - 1)
                if end_ci > ci:
                    cell = cell.merge(row.cells[end_ci])

            # Set cell content
            _fill_cell(cell, cell_data, converter,
                       cell_data.align or columns[ci].align)

            ci += cell_data.colspan

    # Style: make first row bold (header) and force black text on ALL cells.
    # We set BOTH w:color/@val="000000" AND w:color/@themeColor="text1" to
    # prevent any theme / conditional-format override in Word for Mac.
    for ri, row_obj in enumerate(table.rows):
        for cell in row_obj.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if ri == 0:
                        run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    # Also set themeColor to block theme overrides
                    rPr = run._element.get_or_add_rPr()
                    color_el = rPr.find(qn("w:color"))
                    if color_el is not None:
                        color_el.set(qn("w:themeColor"), "text1")


def _apply_cell_margins(table):
    """Set tight cell margins on the table to match LaTeX table spacing."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblCellMar = OxmlElement("w:tblCellMar")
    # top/bottom: 28 twips ≈ 0.5mm — keeps text close to border lines
    # left/right: 57 twips ≈ 1mm — minimal horizontal padding
    for side, val in [("top", "28"), ("bottom", "28"),
                      ("left", "57"), ("right", "57")]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)


def _apply_column_widths(table, columns: list[ColumnDef]):
    """Apply column widths to the table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Set table width to 100%
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # Calculate proportional widths
    total_fixed = sum(c.width_cm for c in columns if c.width_cm)
    num_auto = sum(1 for c in columns if c.width_cm is None)
    page_width_cm = 15.0

    auto_width = 0
    if num_auto > 0:
        remaining = page_width_cm - total_fixed
        auto_width = max(remaining / num_auto, 1.0)

    col_widths_cm = []
    for col in columns:
        if col.width_cm:
            col_widths_cm.append(col.width_cm)
        else:
            col_widths_cm.append(auto_width)

    total_width = sum(col_widths_cm)

    # Set gridCol
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(1 if tblPr is not None else 0, tblGrid)
    for gc in list(tblGrid.findall(qn("w:gridCol"))):
        tblGrid.remove(gc)

    page_tw = 9520  # twips for ~16cm
    for w in col_widths_cm:
        gc = OxmlElement("w:gridCol")
        tw = int(page_tw * w / total_width)
        gc.set(qn("w:w"), str(tw))
        tblGrid.append(gc)

    # Set cell widths
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci >= len(columns):
                break
            tcPr = cell._tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell._tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            pct_val = int(5000 * col_widths_cm[ci] / total_width)
            tcW.set(qn("w:w"), str(pct_val))
            tcW.set(qn("w:type"), "pct")


def _apply_borders(table, border_style: BorderStyle, num_rows: int):
    """Apply border style to the table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))

    # Remove existing style
    tblStyle = tblPr.find(qn("w:tblStyle"))
    if tblStyle is not None:
        tblPr.remove(tblStyle)

    old_borders = tblPr.find(qn("w:tblBorders"))
    if old_borders is not None:
        tblPr.remove(old_borders)

    tblBorders = OxmlElement("w:tblBorders")

    if border_style.style == "three_line":
        THICK = "12"
        NONE = "0"
        for name, val, sz in [
            ("top", "single", THICK),
            ("left", "none", NONE),
            ("bottom", "single", THICK),
            ("right", "none", NONE),
            ("insideH", "none", NONE),
            ("insideV", "none", NONE),
        ]:
            b = OxmlElement(f"w:{name}")
            b.set(qn("w:val"), val)
            b.set(qn("w:sz"), sz)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tblBorders.append(b)

        # Add thin border below header row (first row)
        if num_rows > 0:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell._tc.insert(0, tcPr)
                tcBorders = OxmlElement("w:tcBorders")
                btm = OxmlElement("w:bottom")
                btm.set(qn("w:val"), "single")
                btm.set(qn("w:sz"), "6")  # 0.75pt
                btm.set(qn("w:space"), "0")
                btm.set(qn("w:color"), "000000")
                tcBorders.append(btm)
                tcPr.append(tcBorders)

    elif border_style.style == "hline_only":
        # \hline without | in column spec → horizontal borders only
        # Matches PDF rendering: top/bottom/insideH lines, no vertical
        LINE = "4"
        NONE = "0"
        for name, val, sz in [
            ("top", "single", LINE),
            ("left", "none", NONE),
            ("bottom", "single", LINE),
            ("right", "none", NONE),
            ("insideH", "single", LINE),
            ("insideV", "none", NONE),
        ]:
            b = OxmlElement(f"w:{name}")
            b.set(qn("w:val"), val)
            b.set(qn("w:sz"), sz)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tblBorders.append(b)

    elif border_style.style == "grid":
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{name}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tblBorders.append(b)

    else:  # "none"
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{name}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            tblBorders.append(b)

    tblPr.append(tblBorders)


def _fill_cell(
    cell,
    cell_data: CellData,
    converter: "LatexToDocxConverter",
    align: str = "left",
):
    """Fill a table cell with content from tokens."""
    # Clear default paragraph
    for p in cell.paragraphs:
        if p.text == "":
            pass  # keep the default empty paragraph

    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

    align_map = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    }
    para.alignment = align_map.get(align, WD_PARAGRAPH_ALIGNMENT.LEFT)

    # Override inherited Normal style spacing — table cells need tight layout
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.line_spacing = 1.0

    # Vertical center — matches LaTeX default cell alignment
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._tc.insert(0, tcPr)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)

    # Read fonts from profile
    profile = getattr(converter, "profile", None)
    body_latin = profile.fonts.body_latin if profile else "Times New Roman"
    body_east_asian = profile.fonts.body_east_asian if profile else "STSong"

    # Convert tokens to text with formatting
    tokens = cell_data.tokens
    text = _tokens_to_cell_text(tokens)
    if text.strip():
        # Clear existing runs
        for run in para.runs:
            run.text = ""
        run = para.add_run(text.strip())
        run.font.size = Pt(10.5)
        run.font.name = body_latin
        # Set East Asian font
        if body_east_asian:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), body_east_asian)


def _tokens_to_cell_text(tokens: list[Token]) -> str:
    """Convert cell tokens to plain text."""
    parts = []
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == TokenType.TEXT:
            parts.append(tok.value)
        elif tok.type == TokenType.WHITESPACE:
            parts.append(" ")
        elif tok.type == TokenType.COMMAND:
            name = tok.extra.get("name", "")
            # Symbol commands → Unicode
            if name in SYMBOL_MAP:
                parts.append(SYMBOL_MAP[name])
                i += 1
                continue
            if name in ("textbf", "textit", "emph", "underline",
                        "heiti", "songti", "kaiti", "fangsong",
                        "text", "textrm", "texttt", "textsf"):
                # Consume brace group, extract text
                j = i + 1
                while j < len(tokens) and tokens[j].type == TokenType.WHITESPACE:
                    j += 1
                if j < len(tokens) and tokens[j].type == TokenType.BRACE_OPEN:
                    depth = 1
                    j += 1
                    inner = []
                    while j < len(tokens) and depth > 0:
                        if tokens[j].type == TokenType.BRACE_OPEN:
                            depth += 1
                        elif tokens[j].type == TokenType.BRACE_CLOSE:
                            depth -= 1
                            if depth == 0:
                                j += 1
                                break
                        inner.append(tokens[j])
                        j += 1
                    parts.append(_tokens_to_cell_text(inner))
                    i = j
                    continue
            elif name == "multicolumn":
                # Already handled at row level
                pass
            elif name in ("centering", "raggedright", "raggedleft",
                          "bfseries", "itshape"):
                pass
            elif name == "makecell":
                # \makecell{line1 \\ line2}
                j = i + 1
                while j < len(tokens) and tokens[j].type == TokenType.WHITESPACE:
                    j += 1
                if j < len(tokens) and tokens[j].type == TokenType.BRACE_OPEN:
                    depth = 1
                    j += 1
                    inner = []
                    while j < len(tokens) and depth > 0:
                        if tokens[j].type == TokenType.BRACE_OPEN:
                            depth += 1
                        elif tokens[j].type == TokenType.BRACE_CLOSE:
                            depth -= 1
                            if depth == 0:
                                j += 1
                                break
                        inner.append(tokens[j])
                        j += 1
                    cell_text = _tokens_to_cell_text(inner)
                    parts.append(cell_text)
                    i = j
                    continue
        elif tok.type == TokenType.MATH_INLINE:
            from .math_handler import _latex_math_to_text
            parts.append(_latex_math_to_text(tok.extra.get("content", "")))
        elif tok.type == TokenType.NEWLINE_CMD:
            parts.append("\n")
        i += 1
    return normalize_latex_text("".join(parts))
