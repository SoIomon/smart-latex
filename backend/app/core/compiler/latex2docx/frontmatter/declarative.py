"""Declarative front-matter builder — interprets JSON config from DocxProfile.

Replaces per-template Python classes with a single engine that reads
``profile.frontmatter.sections`` and builds the front-matter elements.
"""

from __future__ import annotations

import re
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.core.compiler.word_preprocessor import WordExportMetadata
from app.core.compiler.word_postprocessor import (
    _make_paragraph as make_paragraph,
    _make_page_break as make_page_break,
    _make_section_break as make_section_break,
    _make_logo_paragraph as make_logo_paragraph,
    _make_info_table as make_info_table,
    _make_toc_field_paragraph as make_toc_field_paragraph,
    _set_static_header as set_static_header,
    _set_styleref_header as set_styleref_header,
)
from . import FrontmatterBuilder

if TYPE_CHECKING:
    from app.core.compiler.latex2docx.profile import (
        DocxProfile,
        FrontmatterElementConfig,
        FrontmatterSectionConfig,
    )

_ALIGN_MAP = {
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
}


def _make_spacing_para(space_pt: float):
    """Create a zero-height paragraph that only contributes vertical spacing.

    Sets line spacing to exactly 0 so the paragraph itself is invisible;
    only ``w:spacing/@w:after`` adds the desired gap.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "0")
    spacing.set(qn("w:lineRule"), "exact")
    twips = int(space_pt * 20)
    spacing.set(qn("w:after"), str(twips))
    pPr.append(spacing)

    # Set font size to 1pt so Word doesn't add extra height
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "2")  # 1pt in half-points
    rPr.append(sz)
    pPr.append(rPr)

    p.append(pPr)
    return p


class DeclarativeFrontmatterBuilder(FrontmatterBuilder):
    """Builds front-matter by interpreting the profile's JSON configuration."""

    def build(self, doc: Document, metadata: WordExportMetadata) -> None:
        self._metadata = metadata  # store for _apply_page_headers
        self._build_frontmatter(doc, metadata)
        self._apply_body_section_breaks(doc)
        self._apply_page_headers(doc)

    def should_handle_command(self, cmd: str) -> bool:
        return cmd in ("maketitle", "MAKETITLE", "makedeclaration",
                       "tableofcontents", "frontmatter", "mainmatter")

    # -- Frontmatter sections --------------------------------------------------

    def _build_frontmatter(self, doc: Document, metadata: WordExportMetadata):
        if not self.profile:
            return

        body = doc.element.body
        first_element = body[0] if len(body) > 0 else None

        elements = []
        for section_cfg in self.profile.frontmatter.sections:
            # Check condition
            if section_cfg.condition:
                if not getattr(metadata, section_cfg.condition, None):
                    continue

            for elem_cfg in section_cfg.elements:
                # Check element-level condition
                if elem_cfg.condition:
                    if not getattr(metadata, elem_cfg.condition, None):
                        continue

                handler = self._ELEMENT_HANDLERS.get(elem_cfg.type)
                if handler:
                    result = handler(self, doc, elem_cfg, metadata)
                    if result is not None:
                        # For non-text elements (logo, info_table, etc.) that
                        # don't natively support space_before, insert a zero-
                        # height spacing paragraph before them.
                        if (elem_cfg.space_before_pt
                                and elem_cfg.type not in ("text", "spacer")):
                            elements.append(
                                _make_spacing_para(elem_cfg.space_before_pt)
                            )
                        if isinstance(result, list):
                            elements.extend(result)
                        else:
                            elements.append(result)

            if section_cfg.break_after:
                elements.append(make_section_break(section_cfg.break_after))

        # Insert all elements at position 0
        if first_element is not None:
            insert_idx = body.index(first_element)
            for elem in elements:
                body.insert(insert_idx, elem)
                insert_idx += 1
        else:
            for elem in elements:
                body.append(elem)

    # -- Element handlers ------------------------------------------------------

    def _build_text(self, doc, cfg: "FrontmatterElementConfig", metadata):
        """Build a text paragraph, resolving {field} placeholders."""
        content = cfg.content
        if cfg.field:
            content = getattr(metadata, cfg.field, "") or ""
        elif content:
            # Resolve {field} placeholders in content
            content = self._interpolate(content, metadata)

        if not content and not cfg.content:
            return None

        return make_paragraph(
            content,
            font_name=cfg.font,
            font_size=Pt(cfg.size_pt),
            bold=cfg.bold,
            alignment=_ALIGN_MAP.get(cfg.align),
            space_before=Pt(cfg.space_before_pt) if cfg.space_before_pt else None,
        )

    def _build_spacer(self, doc, cfg: "FrontmatterElementConfig", metadata):
        """Build one or more blank paragraphs."""
        return [make_paragraph("") for _ in range(cfg.lines)]

    def _build_logo(self, doc, cfg: "FrontmatterElementConfig", metadata):
        """Build the school logo paragraph."""
        logo = make_logo_paragraph(doc, metadata)
        if logo is not None and cfg.space_before_pt:
            # Add spacing to the logo paragraph via w:pPr/w:spacing
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            pPr = logo.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                logo.insert(0, pPr)
            spacing = OxmlElement("w:spacing")
            twips = int(cfg.space_before_pt * 20)  # pt to twips
            spacing.set(qn("w:before"), str(twips))
            pPr.append(spacing)
        return logo

    def _build_info_table(self, doc, cfg: "FrontmatterElementConfig", metadata):
        """Build a 2-column info table with label:value rows."""
        rows = []
        for row_data in cfg.rows:
            if len(row_data) >= 2:
                label = row_data[0]
                value_template = row_data[1]
                value = self._interpolate(value_template, metadata)
                rows.append((label, value))
        if rows:
            return make_info_table(rows)
        return None

    def _build_boilerplate(self, doc, cfg: "FrontmatterElementConfig", metadata):
        """Build multiple boilerplate text paragraphs."""
        elements = []
        first = True
        for row_data in cfg.rows:
            text = row_data[0] if row_data else ""
            text = self._interpolate(text, metadata)
            if text:
                elements.append(make_paragraph(
                    text,
                    font_name=cfg.font,
                    font_size=Pt(cfg.size_pt),
                    bold=cfg.bold,
                    alignment=_ALIGN_MAP.get(cfg.align),
                    space_before=Pt(cfg.space_before_pt) if cfg.space_before_pt and first else None,
                ))
                first = False
        return elements

    def _build_signature_block(self, doc, cfg: "FrontmatterElementConfig", metadata):
        """Build signature line paragraphs."""
        elements = []
        for row_data in cfg.rows:
            text = row_data[0] if row_data else ""
            text = self._interpolate(text, metadata)
            elements.append(make_paragraph(
                text,
                font_name=cfg.font,
                font_size=Pt(cfg.size_pt),
                alignment=_ALIGN_MAP.get(cfg.align),
            ))
        return elements

    _ELEMENT_HANDLERS = {
        "text": _build_text,
        "spacer": _build_spacer,
        "logo": _build_logo,
        "info_table": _build_info_table,
        "boilerplate": _build_boilerplate,
        "signature_block": _build_signature_block,
    }

    # -- Body section breaks ---------------------------------------------------

    def _apply_body_section_breaks(self, doc: Document):
        """Insert section breaks before headings matching the configured rules."""
        if not self.profile:
            return

        breaks = self.profile.frontmatter.body_section_breaks
        if not breaks:
            return

        body = doc.element.body
        matched_patterns: set[str] = set()

        for para in doc.paragraphs:
            # Match both Heading 1 style AND non-Heading paragraphs that
            # look like headings (used for unnumbered headings excluded from TOC)
            is_heading1 = para.style and para.style.style_id == "Heading1"
            is_heading_like = (
                not is_heading1
                and para.runs
                and para.runs[0].bold
                and para.runs[0].font.size
                and para.runs[0].font.size >= Pt(14)
            )
            if not (is_heading1 or is_heading_like):
                continue
            text = para.text.strip()

            for brk in breaks:
                # Skip if first_only and already matched
                brk_key = brk.before_heading_text or brk.before_heading_pattern
                if brk.first_only and brk_key in matched_patterns:
                    continue

                match = False
                if brk.before_heading_text and text == brk.before_heading_text:
                    match = True
                elif brk.before_heading_pattern and re.match(brk.before_heading_pattern, text):
                    match = True

                if match:
                    try:
                        idx = list(body).index(para._element)
                        body.insert(idx, make_section_break(brk.break_type))
                        matched_patterns.add(brk_key)
                    except ValueError:
                        pass
                    break  # only one break per heading

        # Handle auto TOC insertion
        auto_toc = self.profile.frontmatter.auto_toc
        if auto_toc and auto_toc.insert_before_first_chapter:
            self._insert_auto_toc(doc, auto_toc)

    def _insert_auto_toc(self, doc: Document, auto_toc):
        """Insert TOC before the first chapter heading if no TOC heading exists."""
        body = doc.element.body
        first_chapter_elem = None
        toc_exists = False

        for para in doc.paragraphs:
            text = para.text.strip()
            is_heading1 = para.style and para.style.style_id == "Heading1"

            if "目" in text and "录" in text:
                toc_exists = True
            elif is_heading1 and re.match(r"第\s*\d+\s*章", text) and first_chapter_elem is None:
                first_chapter_elem = para._element

        if toc_exists or first_chapter_elem is None:
            return

        try:
            idx = list(body).index(first_chapter_elem)
            toc_elems = [
                make_section_break("oddPage"),
                make_paragraph(
                    auto_toc.heading_text,
                    font_name=auto_toc.heading_font, font_size=Pt(16), bold=True,
                    alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                ),
                make_paragraph(""),
                make_toc_field_paragraph(),
                make_section_break("oddPage"),
            ]
            for i, elem in enumerate(toc_elems):
                body.insert(idx + i, elem)
        except ValueError:
            pass

    # -- Page headers, footers, page numbering -----------------------------------

    def _apply_page_headers(self, doc: Document):
        """Apply page headers, footers, and page numbering.

        Follows LaTeX semantics rather than pattern-matching:
        1. Cover sections (generated by frontmatter.sections) → no header, no page number
        2. Front-matter sections (between cover and first body chapter) →
           static header = first heading-like text in that section, Roman page numbers
        3. Body sections (containing numbered chapter headings) →
           STYLEREF header, Arabic page numbers starting at 1
        """
        if not self.profile:
            return

        ph = self.profile.page_headers
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn

        # Override profile values with auto-detected metadata.
        # metadata fields are None when not detected → fall back to profile.
        # metadata.twoside is False when not detected → fall back to profile.
        metadata = self._metadata
        _meta_fm = getattr(metadata, 'frontmatter_page_format', None)
        _meta_body = getattr(metadata, 'body_page_format', None)
        fm_fmt = _meta_fm if _meta_fm is not None else ph.frontmatter_page_format
        body_fmt = _meta_body if _meta_body is not None else ph.body_page_format
        use_odd_even = metadata.twoside if getattr(metadata, 'twoside', False) else ph.odd_even

        # Enable odd/even headers
        if use_odd_even:
            settings = doc.settings.element
            for old in settings.findall(_qn("w:evenAndOddHeaders")):
                settings.remove(old)
            eaoh = OxmlElement("w:evenAndOddHeaders")
            settings.append(eaoh)

        # Build section→content map
        body = doc.element.body
        section_elements: list[list] = []
        current: list = []
        for child in body:
            current.append(child)
            pPr = child.find(_qn("w:pPr"))
            if pPr is not None and pPr.find(_qn("w:sectPr")) is not None:
                section_elements.append(current)
                current = []
        section_elements.append(current)

        # Count cover sections — each frontmatter.sections entry with
        # break_after creates one Word section.
        cover_count = sum(
            1 for s in self.profile.frontmatter.sections if s.break_after
        )

        # Find the first *numbered chapter* section (body), skipping
        # unnumbered front-matter headings like 摘要, Abstract, 目录.
        # Read from profile.numbering.unnumbered_headings to stay in sync.
        _heading1_ids = {"Heading1", "LaTeXHeading1"}
        _unnumbered = set(self.profile.numbering.unnumbered_headings) if self.profile else set()
        _frontmatter_titles = _unnumbered | {"图形列表", "表格列表"}
        # Normalize whitespace in the title set for consistent matching
        _frontmatter_titles = {re.sub(r"\s+", " ", t) for t in _frontmatter_titles}
        sections = list(doc.sections)
        first_body_idx = len(sections)  # default: no body found
        for si in range(cover_count, len(sections)):
            if si >= len(section_elements):
                continue
            for elem in section_elements[si]:
                if elem.tag != _qn("w:p"):
                    continue
                pPr = elem.find(_qn("w:pPr"))
                if pPr is None:
                    continue
                pStyle = pPr.find(_qn("w:pStyle"))
                if pStyle is None or pStyle.get(_qn("w:val")) not in _heading1_ids:
                    continue
                # Found a Heading1 — check if it's a front-matter title
                texts = [t.text for t in elem.iter(_qn("w:t")) if t.text]
                heading_text = "".join(texts).strip()
                # Normalize whitespace for matching (e.g. "摘  要" → "摘 要")
                normalized = re.sub(r"\s+", " ", heading_text)
                if normalized in _frontmatter_titles:
                    break  # skip this section, it's front-matter
                # This is a real body chapter heading
                first_body_idx = si
                break
            if first_body_idx < len(sections):
                break

        # Get document title for even-page headers
        doc_title = getattr(metadata, 'title', '') if metadata else ''

        first_frontmatter = True
        first_body = True

        for i, section in enumerate(sections):
            if i < cover_count:
                # ── Cover / declaration: no header, no page numbers ──
                self._clear_section_headers_and_footers(section, ph, use_odd_even)

            elif i < first_body_idx:
                # ── Front matter (摘要, Abstract, 目录, LOF, LOT, etc.) ──
                # Header text = first heading-like text in this section
                heading_text = self._find_section_heading(
                    section_elements[i] if i < len(section_elements) else []
                )
                if heading_text:
                    set_static_header(
                        section, heading_text,
                        font_name=ph.header_font,
                        font_size_pt=ph.header_font_size_pt,
                    )
                else:
                    self._clear_header(section)

                if use_odd_even and doc_title:
                    self._set_even_header(section, doc_title, ph)

                self._set_footer_page_number(section, ph, use_odd_even)
                # Set page format on EVERY frontmatter section (some Word
                # versions don't inherit pgNumType from previous sections)
                self._set_page_number_format(
                    section, fm_fmt,
                    start=1 if first_frontmatter else None,
                )
                first_frontmatter = False

            else:
                # ── Body (each chapter in its own section) ──
                # Use the first heading-like text as static header,
                # matching LaTeX's \leftmark behavior per chapter.
                heading_text = self._find_section_heading(
                    section_elements[i] if i < len(section_elements) else []
                )
                if heading_text:
                    set_static_header(
                        section, heading_text,
                        font_name=ph.header_font,
                        font_size_pt=ph.header_font_size_pt,
                    )

                if use_odd_even and doc_title:
                    self._set_even_header(section, doc_title, ph)

                self._set_footer_page_number(section, ph, use_odd_even, body_mode=True)
                if first_body:
                    self._set_page_number_format(
                        section, body_fmt, start=1,
                    )
                    first_body = False

    # -- Header/footer helpers -------------------------------------------------

    @staticmethod
    def _section_has_heading1(elements: list) -> bool:
        """Check if any element in this section uses a Heading 1 style.

        Matches both built-in ``Heading1`` and custom ``LaTeXHeading1``.
        """
        from docx.oxml.ns import qn as _qn
        _heading1_ids = {"Heading1", "LaTeXHeading1"}
        for elem in elements:
            if elem.tag != _qn("w:p"):
                continue
            pPr = elem.find(_qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(_qn("w:pStyle"))
                if pStyle is not None and pStyle.get(_qn("w:val")) in _heading1_ids:
                    return True
        return False

    @staticmethod
    def _find_section_heading(elements: list) -> str:
        """Find the first heading-like text in a list of OxmlElements.

        Checks two cases:
        1. Paragraphs with a Heading style (Heading1, Heading2, etc.)
        2. Paragraphs with direct bold + large-font formatting
           (from ``_add_heading_no_toc``)

        Returns the text, corresponding to LaTeX's ``\\leftmark``.
        """
        from docx.oxml.ns import qn as _qn

        for elem in elements:
            if elem.tag != _qn("w:p"):
                continue

            pPr = elem.find(_qn("w:pPr"))
            texts = [t.text for t in elem.iter(_qn("w:t")) if t.text]
            text = "".join(texts).strip()
            if not text:
                continue

            # Case 1: Heading style (formatting comes from style, not run).
            # Match both built-in "HeadingN" and custom "LaTeXHeadingN".
            if pPr is not None:
                pStyle = pPr.find(_qn("w:pStyle"))
                if pStyle is not None:
                    style_id = pStyle.get(_qn("w:val"), "")
                    if style_id.startswith("Heading") or style_id.startswith("LaTeXHeading"):
                        return text

            # Case 2: Direct bold + large font (from _add_heading_no_toc)
            runs = elem.findall(_qn("w:r"))
            if not runs:
                continue
            first_rPr = runs[0].find(_qn("w:rPr"))
            if first_rPr is None:
                continue
            bold = first_rPr.find(_qn("w:b"))
            if bold is None:
                continue
            sz = first_rPr.find(_qn("w:sz"))
            if sz is None:
                continue
            try:
                half_pts = int(sz.get(_qn("w:val"), "0"))
            except (ValueError, TypeError):
                continue
            if half_pts >= 28:  # >= 14pt
                return text

        return ""

    def _clear_header(self, section):
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()

    def _clear_section_headers_and_footers(self, section, ph, use_odd_even=None):
        """Clear all headers and footers on a section."""
        odd_even = use_odd_even if use_odd_even is not None else ph.odd_even
        self._clear_header(section)
        if odd_even:
            header = section.even_page_header
            header.is_linked_to_previous = False
            for p in header.paragraphs:
                p.clear()
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        if odd_even:
            even_footer = section.even_page_footer
            even_footer.is_linked_to_previous = False
            for p in even_footer.paragraphs:
                p.clear()

    def _set_even_header(self, section, title: str, ph):
        """Set even-page header with document title."""
        set_static_header(
            section, title,
            font_name=ph.header_font,
            font_size_pt=ph.header_font_size_pt,
            even_page=True,
        )

    def _set_footer_page_number(self, section, ph, use_odd_even=None, body_mode=False):
        """Set page number in footer.

        *body_mode*: when True and odd/even is enabled, odd-page footer is
        right-aligned and even-page footer is left-aligned (matching
        LaTeX twoside behaviour).  When False (front-matter), both are centred.
        """
        from app.core.compiler.word_postprocessor import (
            _add_page_field as add_page_field,
        )

        odd_even = use_odd_even if use_odd_even is not None else ph.odd_even

        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()

        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        if body_mode and odd_even:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        add_page_field(para)

        # Even page footer
        if odd_even:
            even_footer = section.even_page_footer
            even_footer.is_linked_to_previous = False
            for p in even_footer.paragraphs:
                p.clear()
            if even_footer.paragraphs:
                epara = even_footer.paragraphs[0]
            else:
                epara = even_footer.add_paragraph()
            if body_mode:
                epara.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                epara.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            add_page_field(epara)

    def _set_page_number_format(self, section, fmt: str, start: int | None = 1):
        """Set page number format on a section.

        *start*: page number to start at, or None to continue from previous.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn

        sectPr = section._sectPr
        for old in sectPr.findall(_qn("w:pgNumType")):
            sectPr.remove(old)
        pgNumType = OxmlElement("w:pgNumType")
        pgNumType.set(_qn("w:fmt"), fmt)
        if start is not None:
            pgNumType.set(_qn("w:start"), str(start))
        sectPr.append(pgNumType)

    # -- Helpers ---------------------------------------------------------------

    @staticmethod
    def _interpolate(template: str, metadata) -> str:
        """Replace {field_name} placeholders with metadata values."""
        def replacer(match):
            field_name = match.group(1)
            return str(getattr(metadata, field_name, "") or "")
        return re.sub(r"\{(\w+)\}", replacer, template)
