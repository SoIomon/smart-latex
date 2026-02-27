"""Layer 4: python-docx post-processor for Word export.

Fixes styles, page layout, page numbers, and optionally rebuilds the
cover page and revision-records table that were stripped during
preprocessing.
"""

import logging
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor, Emu

from app.core.compiler.word_preprocessor import WordExportMetadata
from app.core.fonts import get_cjk_fonts

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def postprocess_word(
    docx_path: str | Path,
    metadata: WordExportMetadata,
    template_id: str = "",
) -> None:
    """Apply generic post-processing fixes to a Word document."""
    doc = Document(str(docx_path))

    # Load profile for template-specific settings
    from app.core.compiler.latex2docx.profile import load_profile
    profile = load_profile(template_id)

    # Phase 1: content & style fixes (before section breaks are created)
    _fix_styles(doc, profile)
    _fix_list_bullets(doc)
    _fix_chapter_headings(doc, profile)

    # Phase 2: structural changes that create new sections.
    # Template-specific front-matter is now handled by latex2docx builders.
    if metadata.has_cover:
        _rebuild_cover_page(doc, metadata)

    if metadata.revision_records:
        _rebuild_revision_table(doc, metadata)

    # Phase 3: layout & formatting (after all sections exist)
    _fix_page_layout(doc, metadata)
    _fix_table_widths(doc)
    _add_page_numbers(doc)

    # Tell Word to auto-update all fields (TOC, STYLEREF, PAGE) on open
    _enable_update_fields(doc)

    doc.save(str(docx_path))


# ---------------------------------------------------------------------------
# Page layout
# ---------------------------------------------------------------------------

def _fix_page_layout(doc: Document, metadata: WordExportMetadata) -> None:
    """Set page margins from extracted geometry metadata."""
    geo = metadata.geometry
    if not geo:
        return

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.orientation = WD_ORIENT.PORTRAIT

        if "top" in geo:
            section.top_margin = _parse_length(geo["top"])
        if "bottom" in geo:
            section.bottom_margin = _parse_length(geo["bottom"])
        if "left" in geo:
            section.left_margin = _parse_length(geo["left"])
        if "right" in geo:
            section.right_margin = _parse_length(geo["right"])


def _parse_length(value: str) -> int:
    """Convert a LaTeX length string (e.g. '2.54cm') to EMU."""
    import re
    m = re.match(r"([\d.]+)\s*(cm|mm|in|pt|bp)", value.strip())
    if not m:
        return Cm(2.54)  # fallback
    num = float(m.group(1))
    unit = m.group(2)
    if unit == "cm":
        return Cm(num)
    elif unit == "mm":
        return Emu(int(num * 36000))
    elif unit == "in":
        return Emu(int(num * 914400))
    elif unit in ("pt", "bp"):
        return Pt(num)
    return Cm(num)


# ---------------------------------------------------------------------------
# Style fixes
# ---------------------------------------------------------------------------

def _fix_styles(doc: Document, profile=None) -> None:
    """Ensure correct fonts on Normal and Heading styles.

    If *profile* is provided, fonts/sizes/indent are read from it;
    otherwise the hardcoded defaults are used for backward compatibility.
    """
    # Import here to avoid circular imports at module level
    if profile is None:
        from app.core.compiler.latex2docx.profile import DocxProfile
        profile = DocxProfile()

    f = profile.fonts
    s = profile.styles

    # Normal style
    try:
        normal = doc.styles["Normal"]
        normal.font.name = f.body_latin
        normal.font.size = Pt(s.normal.font_size_pt)
        normal.paragraph_format.first_line_indent = Pt(s.normal.first_line_indent_pt)
        if f.body_east_asian:
            _set_east_asian_font(normal, f.body_east_asian)
    except KeyError:
        pass

    # Heading styles
    for hs in s.headings:
        try:
            style = doc.styles[f"Heading {hs.level}"]
            style.font.name = f.heading_latin
            style.font.size = Pt(hs.font_size_pt)
            style.font.bold = hs.bold
            style.font.color.rgb = RGBColor(0, 0, 0)
            if f.heading_east_asian:
                _set_east_asian_font(style, f.heading_east_asian)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.first_line_indent = Pt(0)  # no indent on headings

            # Remove any numPr from heading style.
            pPr = style.element.get_or_add_pPr()
            for old_numPr in pPr.findall(qn("w:numPr")):
                pPr.remove(old_numPr)
        except KeyError:
            pass

    # Caption style — override the default blue (accent1) with black
    try:
        caption = doc.styles["Caption"]
        caption.font.name = f.heading_latin
        caption.font.size = Pt(s.caption.font_size_pt)
        caption.font.color.rgb = RGBColor(0, 0, 0)
        caption.font.bold = True
        if f.caption_east_asian:
            _set_east_asian_font(caption, f.caption_east_asian)
        caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        caption.paragraph_format.first_line_indent = Pt(0)
    except KeyError:
        pass

    # Title style (in case Pandoc generates a title block)
    try:
        title_style = doc.styles["Title"]
        title_font = f.heading_east_asian or f.heading_latin
        title_style.font.name = title_font
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        if f.heading_east_asian:
            _set_east_asian_font(title_style, f.heading_east_asian)
    except KeyError:
        pass

    # Final sweep: remove numPr from ALL heading paragraphs in the document.
    # Python-docx's default template may leave stale numPr references that
    # cause bullet dots in some Word versions.
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                for numPr in pPr.findall(qn("w:numPr")):
                    pPr.remove(numPr)

    # Keep numbering definitions intact to preserve list semantics.
    # Heading-related phantom bullets are handled by removing heading numPr
    # instead of deleting the entire numbering subsystem.


def _set_east_asian_font(style, font_name: str) -> None:
    """Set the East Asian font on a style element."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


# ---------------------------------------------------------------------------
# Page numbers
# ---------------------------------------------------------------------------

def _add_page_numbers(doc: Document) -> None:
    """Add centered page numbers to the footer of the first section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear existing content
        for p in footer.paragraphs:
            p.clear()

        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _add_page_field(para)


def _enable_update_fields(doc: Document) -> None:
    """Set updateFields in document settings so Word refreshes TOC etc. on open.

    This is generic — benefits any template that uses field codes (TOC, PAGE,
    STYLEREF, etc.).
    """
    settings_elem = doc.settings.element
    # Remove existing updateFields if present
    for uf in settings_elem.findall(qn("w:updateFields")):
        settings_elem.remove(uf)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings_elem.append(update)


def _add_page_field(paragraph) -> None:
    """Insert a PAGE field code into a paragraph."""
    run = paragraph.add_run()
    run.font.size = Pt(10)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._element.append(instrText)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar_end)


# ---------------------------------------------------------------------------
# Cover page rebuild
# ---------------------------------------------------------------------------

def _rebuild_cover_page(doc: Document, metadata: WordExportMetadata) -> None:
    """Insert a cover page at the very beginning of the document."""
    # We insert elements *before* the first existing paragraph.
    # Strategy: collect cover elements, then insert them at position 0.

    body = doc.element.body
    first_element = body[0] if len(body) > 0 else None

    cover_elements = []

    # ── Metadata row: doc_number / phase_mark / classification ──────────
    meta_lines = []
    if metadata.doc_number:
        meta_lines.append(f"文件编号: {metadata.doc_number}")
    if metadata.phase_mark:
        meta_lines.append(f"阶段标志: {metadata.phase_mark}")
    if metadata.classification:
        meta_lines.append(f"密级: {metadata.classification}")

    _cjk = get_cjk_fonts()
    for line in meta_lines:
        p = _make_paragraph(line, font_name=_cjk.heiti, font_size=Pt(10.5))
        cover_elements.append(p)

    # ── Spacer ──────────────────────────────────────────────────────────
    cover_elements.append(_make_paragraph(""))

    # ── Title ───────────────────────────────────────────────────────────
    if metadata.title:
        p = _make_paragraph(
            metadata.title,
            font_name=_cjk.heiti,
            font_size=Pt(22),
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        cover_elements.append(p)
        cover_elements.append(_make_paragraph(""))

    # ── Approval table ──────────────────────────────────────────────────
    approval_rows = [
        ("编写", metadata.writer, metadata.write_date),
        ("校对", metadata.proofreader, metadata.proofread_date),
        ("审核", metadata.reviewer, metadata.review_date),
        ("标审", metadata.standard_reviewer, metadata.standard_review_date),
        ("批准", metadata.approver, metadata.approve_date),
    ]

    # Only add the table if there's at least one non-empty entry
    has_approval = any(name or date for _, name, date in approval_rows)
    if has_approval:
        tbl = _make_approval_table(doc, approval_rows)
        cover_elements.append(tbl)
        cover_elements.append(_make_paragraph(""))

    # ── Institute + date ────────────────────────────────────────────────
    if metadata.institute:
        p = _make_paragraph(
            metadata.institute,
            font_name=_cjk.heiti,
            font_size=Pt(18),
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        cover_elements.append(p)

    if metadata.report_date:
        p = _make_paragraph(
            metadata.report_date,
            font_name=_cjk.heiti,
            font_size=Pt(16),
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        cover_elements.append(p)

    # ── Page break ──────────────────────────────────────────────────────
    cover_elements.append(_make_page_break())

    # Insert all cover elements before the first existing element
    if first_element is not None:
        insert_idx = body.index(first_element)
        for elem in cover_elements:
            body.insert(insert_idx, elem)
            insert_idx += 1
    else:
        for elem in cover_elements:
            body.append(elem)


def _make_paragraph(
    text: str,
    font_name: str = "",
    font_size=Pt(12),
    bold: bool = False,
    alignment=None,
    space_before=None,
    space_after=None,
) -> OxmlElement:
    """Create a w:p element with a single run.

    *space_before* / *space_after*: vertical spacing in Pt (e.g. ``Pt(50)``).
    """
    if not font_name:
        font_name = get_cjk_fonts().songti
    p = OxmlElement("w:p")

    need_pPr = alignment is not None or space_before is not None or space_after is not None
    if need_pPr:
        pPr = OxmlElement("w:pPr")
        if space_before is not None or space_after is not None:
            spacing = OxmlElement("w:spacing")
            if space_before is not None:
                # Pt() returns EMU; OOXML spacing is in twips (1 pt = 20 twips)
                twips = int(space_before / 12700 * 20)
                spacing.set(qn("w:before"), str(twips))
            if space_after is not None:
                twips = int(space_after / 12700 * 20)
                spacing.set(qn("w:after"), str(twips))
            pPr.append(spacing)
        if alignment is not None:
            jc = OxmlElement("w:jc")
            align_map = {
                WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
                WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
                WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            }
            jc.set(qn("w:val"), align_map.get(alignment, "left"))
            pPr.append(jc)
        p.append(pPr)

    if not text:
        return p

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.append(rFonts)

    # Size
    if font_size:
        sz = OxmlElement("w:sz")
        # python-docx Pt returns EMU; Word uses half-points
        half_points = int(font_size / 6350)
        sz.set(qn("w:val"), str(half_points))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(half_points))
        rPr.append(szCs)

    # Bold
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)

    p.append(r)
    return p


def _make_page_break() -> OxmlElement:
    """Create a paragraph containing a page break."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _make_section_break(break_type: str = "oddPage") -> OxmlElement:
    """Create a paragraph with a section break.

    *break_type*: ``oddPage`` | ``evenPage`` | ``nextPage`` | ``continuous``
    ``oddPage`` forces the next content onto an odd (recto) page, inserting a
    blank page if necessary — matching LaTeX ``\\cleardoublepage`` behaviour.

    Includes ``pgSz`` (A4) and ``pgMar`` so Word can render blank pages
    correctly even before ``_fix_page_layout`` runs.
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")

    # Page size — A4 (210mm × 297mm) in twips
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "11906")
    pgSz.set(qn("w:h"), "16838")
    sectPr.append(pgSz)

    # Page margins — defaults matching ucas_thesis geometry (≈2.54 cm)
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"), "1440")     # 1 inch = 1440 twips
    pgMar.set(qn("w:right"), "1440")
    pgMar.set(qn("w:bottom"), "1440")
    pgMar.set(qn("w:left"), "1440")
    pgMar.set(qn("w:header"), "720")
    pgMar.set(qn("w:footer"), "720")
    pgMar.set(qn("w:gutter"), "0")
    sectPr.append(pgMar)

    pgSzType = OxmlElement("w:type")
    pgSzType.set(qn("w:val"), break_type)
    sectPr.append(pgSzType)

    pPr.append(sectPr)
    p.append(pPr)
    return p


def _make_logo_paragraph(doc: Document, metadata):
    """Create a centered paragraph with the school logo image, or None.

    The image is inserted into *doc* so that the relationship (r:embed rId)
    lives in the real document package.  The paragraph element is then detached
    from its current position and returned for re-insertion at the desired
    location.
    """
    if not metadata.school_logo:
        return None

    from app.core.templates.registry import get_template_dir
    template_dir = get_template_dir(metadata.template_id)
    if not template_dir:
        return None

    # Find the logo file — could be .pdf, .png, .jpg etc.
    logo_name = metadata.school_logo
    img_dir = template_dir / "Img"
    logo_path = None
    for ext in ("", ".pdf", ".png", ".jpg", ".jpeg", ".eps"):
        candidate = img_dir / f"{logo_name}{ext}"
        if candidate.exists():
            logo_path = candidate
            break

    if logo_path is None:
        return None

    # If PDF, convert first page to PNG using PyMuPDF
    # Also compute the scaled width to match the LaTeX rendering
    logo_width_cm = None
    if logo_path.suffix.lower() == ".pdf":
        logo_width_cm = _get_pdf_scaled_width(logo_path, metadata.school_logo_scale)
        logo_path = _pdf_to_png(logo_path)
        if logo_path is None:
            return None

    if logo_width_cm is None:
        logo_width_cm = 10.0  # fallback

    # Insert the picture into the real document so the image relationship
    # (r:embed rId) is valid.  add_picture() appends a new paragraph at the
    # end of the document body; we grab it, detach it, and return the element
    # for the caller to place wherever it likes.
    try:
        doc.add_picture(str(logo_path), width=Cm(logo_width_cm))
        pic_para = doc.paragraphs[-1]
        pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        body = doc.element.body
        body.remove(pic_para._element)
        return pic_para._element
    except Exception as e:
        logger.warning("Failed to insert school logo: %s", e)
        return None


def _get_pdf_scaled_width(pdf_path: Path, scale: float):
    """Return the width in cm of a PDF page scaled by *scale*."""
    try:
        import fitz
        doc = fitz.open(str(pdf_path))
        width_pt = doc[0].rect.width
        doc.close()
        if scale > 0:
            return width_pt * scale / 72 * 2.54  # pt → cm
    except Exception:
        pass
    return None


def _pdf_to_png(pdf_path: Path):
    """Convert the first page of a PDF to a temporary PNG file."""
    try:
        import fitz  # PyMuPDF
        import tempfile
        doc = fitz.open(str(pdf_path))
        page = doc[0]
        # Render at 3x for good quality
        pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
        tmp = Path(tempfile.mktemp(suffix=".png"))
        pix.save(str(tmp))
        doc.close()
        return tmp
    except Exception as e:
        logger.warning("Failed to convert PDF logo to PNG: %s", e)
        return None


def _make_approval_table(doc: Document, rows: list[tuple[str, str, str]]) -> OxmlElement:
    """Build a simple 3-column approval table as OxmlElement."""
    tbl = OxmlElement("w:tbl")

    # Table properties
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)

    # Table borders
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    tbl.append(tblPr)

    # Header row
    header_row = _make_table_row(["项目", "人员", "日期"], bold=True, font_name=get_cjk_fonts().heiti)
    tbl.append(header_row)

    # Data rows
    for label, name, date in rows:
        row = _make_table_row([label, name, date])
        tbl.append(row)

    return tbl


def _make_table_row(
    cells: list[str], bold: bool = False, font_name: str = ""
) -> OxmlElement:
    """Create a w:tr element with the given cell texts."""
    if not font_name:
        font_name = get_cjk_fonts().songti
    tr = OxmlElement("w:tr")
    for cell_text in cells:
        tc = OxmlElement("w:tc")
        p = OxmlElement("w:p")

        # Center alignment
        pPr = OxmlElement("w:pPr")
        pJc = OxmlElement("w:jc")
        pJc.set(qn("w:val"), "center")
        pPr.append(pJc)
        p.append(pPr)

        if cell_text:
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")

            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)
            rFonts.set(qn("w:eastAsia"), font_name)
            rPr.append(rFonts)

            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "21")  # 10.5pt
            rPr.append(sz)

            if bold:
                b = OxmlElement("w:b")
                rPr.append(b)

            r.append(rPr)

            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = cell_text
            r.append(t)

            p.append(r)

        tc.append(p)
        tr.append(tc)
    return tr


# ---------------------------------------------------------------------------
# Revision table rebuild
# ---------------------------------------------------------------------------

def _rebuild_revision_table(doc: Document, metadata: WordExportMetadata) -> None:
    """Find the simplified revision section and replace it with a proper table.

    If the preprocessor inserted a ``\\section*{文档修改记录}`` that Pandoc
    converted into a heading, we locate it and insert a formatted table
    right after.  If not found, prepend the table after the cover page.
    """
    records = metadata.revision_records
    if not records:
        return

    # Try to locate the revision heading by scanning paragraphs
    insert_after_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if "文档修改记录" in para.text:
            insert_after_idx = idx
            break

    if insert_after_idx is None:
        return  # not found; skip

    # Check if the next element is already a table (Pandoc may have converted it)
    # If so, we leave it as-is (Pandoc did a decent job)
    body = doc.element.body
    ref_element = doc.paragraphs[insert_after_idx]._element

    # Find position in body
    ref_pos = None
    for i, child in enumerate(body):
        if child is ref_element:
            ref_pos = i
            break

    if ref_pos is None:
        return

    # Check next element
    next_pos = ref_pos + 1
    if next_pos < len(body):
        next_elem = body[next_pos]
        # If it's already a table, remove it so we can insert a better one
        if next_elem.tag == qn("w:tbl"):
            body.remove(next_elem)

    # Build the new table
    tbl = _make_revision_table(records)

    # Insert after the heading
    if ref_pos + 1 < len(body):
        body.insert(ref_pos + 1, tbl)
    else:
        body.append(tbl)


def _make_revision_table(records: list[dict]) -> OxmlElement:
    """Create a formatted revision records table."""
    tbl = OxmlElement("w:tbl")

    # Table properties
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    # Borders
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    tbl.append(tblPr)

    # Header
    header = _make_table_row(
        ["版本", "日期", "更改摘要", "修改章节", "备注"],
        bold=True,
        font_name=get_cjk_fonts().heiti,
    )
    tbl.append(header)

    # Data rows
    for rec in records:
        row = _make_table_row([
            rec.get("version", ""),
            rec.get("date", ""),
            rec.get("change_summary", ""),
            rec.get("modified_sections", ""),
            rec.get("remarks", ""),
        ])
        tbl.append(row)

    return tbl


def _make_toc_field_paragraph(
    hint_text: str = "请右键点击此处，选择\u201c更新域\u201d以生成目录",
) -> OxmlElement:
    """Create a paragraph containing a Word TOC field code."""
    toc_para = OxmlElement("w:p")

    run = OxmlElement("w:r")
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run.append(fldChar_begin)
    toc_para.append(run)

    run2 = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-4" \\h \\z \\u '
    run2.append(instrText)
    toc_para.append(run2)

    run3 = OxmlElement("w:r")
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run3.append(fldChar_separate)
    toc_para.append(run3)

    run4 = OxmlElement("w:r")
    rPr4 = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "808080")
    rPr4.append(color)
    run4.append(rPr4)
    t = OxmlElement("w:t")
    t.text = hint_text
    run4.append(t)
    toc_para.append(run4)

    run5 = OxmlElement("w:r")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run5.append(fldChar_end)
    toc_para.append(run5)

    return toc_para


def _make_list_field_paragraph(kind: str = "figure",
                               label: str | None = None) -> OxmlElement:
    """Create a paragraph containing a Word Table of Figures/Tables field.

    *kind*: ``"figure"`` or ``"table"``.
    *label*: The caption label identifier used in ``SEQ`` fields (e.g. "图",
             "表", "Figure").  If *None*, defaults to "图" / "表".

    The field code ``TOC \\c "LABEL"`` tells Word to build a list from
    paragraphs whose captions contain a matching ``SEQ LABEL`` field.
    """
    if label is None:
        label = "图" if kind == "figure" else "表"

    para = OxmlElement("w:p")

    # fldChar begin
    r1 = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fc_begin)
    para.append(r1)

    # instrText
    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\h \\z \\c "{label}" '
    r2.append(instr)
    para.append(r2)

    # fldChar separate
    r3 = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fc_sep)
    para.append(r3)

    # placeholder text
    r4 = OxmlElement("w:r")
    rPr4 = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "808080")
    rPr4.append(color)
    r4.append(rPr4)
    t = OxmlElement("w:t")
    t.text = "请更新域以生成列表"
    r4.append(t)
    para.append(r4)

    # fldChar end
    r5 = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r5.append(fc_end)
    para.append(r5)

    return para


def _make_info_table(rows: list[tuple[str, str]]) -> OxmlElement:
    """Build a 2-column borderless info table (label: value with underline)."""
    _cjk_info = get_cjk_fonts()
    tbl = OxmlElement("w:tbl")

    # Table properties — no borders, centered
    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "4200")   # ~84% of page width — matches LaTeX layout
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)

    # Remove default cell margins so label and value are closer together
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in [("left", "57"), ("right", "57")]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)      # 57 twips ≈ 1mm (default is 108)
        m.set(qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    # Explicitly set no borders
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    tbl.append(tblPr)

    for label, value in rows:
        tr = OxmlElement("w:tr")

        # Label cell — narrow, right-aligned
        tc_label = OxmlElement("w:tc")
        tcPr_label = OxmlElement("w:tcPr")
        tcW_label = OxmlElement("w:tcW")
        tcW_label.set(qn("w:w"), "1800")   # ~36% of table
        tcW_label.set(qn("w:type"), "pct")
        tcPr_label.append(tcW_label)
        # Bottom-align label to match value cell
        vAlign_label = OxmlElement("w:vAlign")
        vAlign_label.set(qn("w:val"), "bottom")
        tcPr_label.append(vAlign_label)
        tc_label.append(tcPr_label)

        p_label = _make_paragraph(
            f"{label}：", font_name=_cjk_info.songti, font_size=Pt(14),
            alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT,
        )
        # Zero spacing on label paragraph too
        p_label_pPr = p_label.find(qn("w:pPr"))
        if p_label_pPr is None:
            p_label_pPr = OxmlElement("w:pPr")
            p_label.insert(0, p_label_pPr)
        p_label_sp = OxmlElement("w:spacing")
        p_label_sp.set(qn("w:before"), "0")
        p_label_sp.set(qn("w:after"), "0")
        p_label_sp.set(qn("w:line"), "240")
        p_label_sp.set(qn("w:lineRule"), "auto")
        p_label_pPr.append(p_label_sp)

        tc_label.append(p_label)
        tr.append(tc_label)

        # Value cell (with bottom border for underline effect)
        tc_val = OxmlElement("w:tc")
        tcPr_val = OxmlElement("w:tcPr")
        tcW_val = OxmlElement("w:tcW")
        tcW_val.set(qn("w:w"), "3200")   # ~64% of table
        tcW_val.set(qn("w:type"), "pct")
        tcPr_val.append(tcW_val)

        tcBorders = OxmlElement("w:tcBorders")
        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "4")
        bottom_border.set(qn("w:space"), "0")
        bottom_border.set(qn("w:color"), "000000")
        tcBorders.append(bottom_border)
        tcPr_val.append(tcBorders)

        # Zero bottom margin so text sits right on the underline
        tcMar = OxmlElement("w:tcMar")
        btm_mar = OxmlElement("w:bottom")
        btm_mar.set(qn("w:w"), "0")
        btm_mar.set(qn("w:type"), "dxa")
        tcMar.append(btm_mar)
        tcPr_val.append(tcMar)

        # Bottom-align text so it sits on the underline border
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "bottom")
        tcPr_val.append(vAlign)

        tc_val.append(tcPr_val)
        p_val = _make_paragraph(
            value, font_name=_cjk_info.songti, font_size=Pt(14),
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        # Remove inherited spacing so text sits tight against the border
        p_val_pPr = p_val.find(qn("w:pPr"))
        if p_val_pPr is None:
            p_val_pPr = OxmlElement("w:pPr")
            p_val.insert(0, p_val_pPr)
        p_spacing = OxmlElement("w:spacing")
        p_spacing.set(qn("w:before"), "0")
        p_spacing.set(qn("w:after"), "0")
        p_spacing.set(qn("w:line"), "240")
        p_spacing.set(qn("w:lineRule"), "auto")
        p_val_pPr.append(p_spacing)
        tc_val.append(p_val)
        tr.append(tc_val)

        tbl.append(tr)

    return tbl


# ---------------------------------------------------------------------------
# Chapter heading format fix
# ---------------------------------------------------------------------------

def _fix_chapter_headings(doc: Document, profile=None) -> None:
    """Convert '1  绪论' → '第 1 章  绪论' for Heading 1 paragraphs."""
    import re as _re

    if profile is None:
        from app.core.compiler.latex2docx.profile import DocxProfile
        profile = DocxProfile()

    skip_set = set(profile.numbering.unnumbered_headings)

    for para in doc.paragraphs:
        if para.style and para.style.style_id == "Heading1":
            text = para.text.strip()
            if text in skip_set:
                continue
            m = _re.match(r"^(\d+)\s+(.+)", text)
            if m:
                num_str, rest = m.group(1), m.group(2)
                new_text = profile.format_chapter(int(num_str), rest)
                if para.runs:
                    for i, run in enumerate(para.runs):
                        if i == 0:
                            run.text = new_text
                        else:
                            run.text = ""
                else:
                    para.text = new_text


# ---------------------------------------------------------------------------
# List bullet fix
# ---------------------------------------------------------------------------

def _fix_list_bullets(doc: Document) -> None:
    """Fix empty bullet characters in numbering definitions."""
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return
    if numbering_part is None:
        return

    numbering_xml = numbering_part.element
    for abstract_num in numbering_xml.iter(qn("w:abstractNum")):
        for lvl in abstract_num.iter(qn("w:lvl")):
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is not None and num_fmt.get(qn("w:val")) == "bullet":
                lvl_text = lvl.find(qn("w:lvlText"))
                if lvl_text is not None:
                    val = lvl_text.get(qn("w:val")) or ""
                    if not val.strip():
                        lvl_text.set(qn("w:val"), "\u2022")
                        rPr = lvl.find(qn("w:rPr"))
                        if rPr is None:
                            rPr = OxmlElement("w:rPr")
                            lvl.append(rPr)
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            rFonts = OxmlElement("w:rFonts")
                            rPr.append(rFonts)
                        rFonts.set(qn("w:ascii"), "Arial Unicode MS")
                        rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")


# ---------------------------------------------------------------------------
# Table width fix
# ---------------------------------------------------------------------------

def _fix_table_widths(doc: Document) -> None:
    """Fix Pandoc tables: 100% width, proportional columns, 三线表 borders."""
    for table in doc.tables:
        tbl = table._tbl

        # Fix table width → 100% (pct = 5000)
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        # Remove tblStyle entirely (we set borders explicitly for 三线表)
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is not None:
            tblPr.remove(tblStyle)

        # ── 三线表 (three-line table) borders ──
        # Table-level: top=thick, bottom=thick, others=none
        old_borders = tblPr.find(qn("w:tblBorders"))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblBorders = OxmlElement("w:tblBorders")
        _THICK = "12"   # 1.5pt in eighth-points
        _NONE = "0"
        for name, val, sz in [
            ("top", "single", _THICK),
            ("left", "none", _NONE),
            ("bottom", "single", _THICK),
            ("right", "none", _NONE),
            ("insideH", "none", _NONE),
            ("insideV", "none", _NONE),
        ]:
            b = OxmlElement(f"w:{name}")
            b.set(qn("w:val"), val)
            b.set(qn("w:sz"), sz)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        # Build proportional column widths from content length
        rows = table.rows
        if not rows:
            continue
        num_cols = len(rows[0].cells)
        if num_cols == 0:
            continue

        # Estimate column widths by max content length
        col_weights = [1] * num_cols
        for row in rows:
            for ci, cell in enumerate(row.cells):
                if ci < num_cols:
                    text_len = len(cell.text.strip())
                    col_weights[ci] = max(col_weights[ci], text_len + 1)

        total_weight = sum(col_weights)
        total_pct = 5000

        # Set gridCol widths in tblGrid
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is None:
            tblGrid = OxmlElement("w:tblGrid")
            tbl.insert(1 if tblPr is not None else 0, tblGrid)
        for gc in list(tblGrid.findall(qn("w:gridCol"))):
            tblGrid.remove(gc)
        page_tw = 9520
        for ci in range(num_cols):
            gc = OxmlElement("w:gridCol")
            tw = int(page_tw * col_weights[ci] / total_weight)
            gc.set(qn("w:w"), str(tw))
            tblGrid.append(gc)

        # Set cell widths (tcW) + header row bottom border
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row.cells):
                if ci >= num_cols:
                    break
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell._tc.insert(0, tcPr)
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.insert(0, tcW)
                pct_val = int(total_pct * col_weights[ci] / total_weight)
                tcW.set(qn("w:w"), str(pct_val))
                tcW.set(qn("w:type"), "pct")

                # First row (header): add thin bottom border (三线表 中线)
                if ri == 0:
                    tcBorders = tcPr.find(qn("w:tcBorders"))
                    if tcBorders is None:
                        tcBorders = OxmlElement("w:tcBorders")
                        tcPr.append(tcBorders)
                    btm = OxmlElement("w:bottom")
                    btm.set(qn("w:val"), "single")
                    btm.set(qn("w:sz"), "6")   # 0.75pt
                    btm.set(qn("w:space"), "0")
                    btm.set(qn("w:color"), "000000")
                    tcBorders.append(btm)


def _set_static_header(section, text: str,
                       font_name: str = "",
                       font_size_pt: float = 10.5,
                       even_page: bool = False) -> None:
    """Set a centered static text header on *section*.

    If *even_page* is True, sets the even-page header instead of the
    default (odd-page) header.
    """
    if not font_name:
        font_name = get_cjk_fonts().songti
    header = section.even_page_header if even_page else section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    if header.paragraphs:
        para = header.paragraphs[0]
    else:
        para = header.add_paragraph()

    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = para.add_run(text)
    run.font.size = Pt(font_size_pt)
    run.font.name = font_name
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)

    # Add bottom border (header rule)
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_styleref_header(section,
                         font_name: str = "",
                         font_size_pt: float = 10.5) -> None:
    """Set a header with STYLEREF field that displays the current Heading 1.

    Each fldChar (begin/separate/end) must be in its own ``<w:r>`` element
    for Word to evaluate the field correctly.
    """
    if not font_name:
        font_name = get_cjk_fonts().songti
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    if header.paragraphs:
        para = header.paragraphs[0]
    else:
        para = header.add_paragraph()

    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    half_points = str(int(font_size_pt * 2))

    def _make_run_rPr():
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), half_points)
        rPr.append(sz)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
        rPr.append(rFonts)
        return rPr

    # Run 1: fldChar begin
    r1 = OxmlElement("w:r")
    r1.append(_make_run_rPr())
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fc_begin)
    para._element.append(r1)

    # Run 2: instrText
    r2 = OxmlElement("w:r")
    r2.append(_make_run_rPr())
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' STYLEREF "heading 1" \\* MERGEFORMAT '
    r2.append(instrText)
    para._element.append(r2)

    # Run 3: fldChar separate
    r3 = OxmlElement("w:r")
    r3.append(_make_run_rPr())
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fc_sep)
    para._element.append(r3)

    # Run 4: placeholder (display text before field update)
    r4 = OxmlElement("w:r")
    r4.append(_make_run_rPr())
    placeholder = OxmlElement("w:t")
    placeholder.text = "章节标题"
    r4.append(placeholder)
    para._element.append(r4)

    # Run 5: fldChar end
    r5 = OxmlElement("w:r")
    r5.append(_make_run_rPr())
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r5.append(fc_end)
    para._element.append(r5)

    # Add bottom border (header rule)
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
