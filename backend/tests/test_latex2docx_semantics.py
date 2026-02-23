from pathlib import Path
import zipfile
import xml.etree.ElementTree as ET
from unittest.mock import patch

from app.core.compiler.latex2docx import convert_latex_to_docx
from app.core.compiler.latex2docx.tex_auxfiles import parse_aux_file
from docx import Document
from app.core.compiler.latex2docx.profile import DocxProfile, LabelsConfig
from app.core.compiler.latex2docx.frontmatter.ucas_thesis import UcasThesisFrontmatter
from app.core.compiler.word_preprocessor import WordExportMetadata


def test_parse_aux_file_uses_bbl_order_for_biblatex_citations(tmp_path: Path):
    aux = tmp_path / "document.aux"
    bbl = tmp_path / "document.bbl"

    aux.write_text(
        "\n".join(
            [
                r"\relax",
                r"\abx@aux@cite{0}{zeta2020}",
                r"\abx@aux@cite{0}{alpha2019}",
            ]
        ),
        encoding="utf-8",
    )
    bbl.write_text(
        "\n".join(
            [
                r"\entry{alpha2019}{article}{}",
                r"\entry{zeta2020}{article}{}",
            ]
        ),
        encoding="utf-8",
    )

    structure = parse_aux_file(aux, bbl_path=bbl)
    assert structure is not None
    assert structure.resolve_ref("alpha2019") == "1"
    assert structure.resolve_ref("zeta2020") == "2"


def test_parse_aux_file_uses_bibitem_order_when_no_bibcite(tmp_path: Path):
    aux = tmp_path / "document.aux"
    bbl = tmp_path / "document.bbl"

    aux.write_text(r"\relax", encoding="utf-8")
    bbl.write_text(
        "\n".join(
            [
                r"\bibitem{beta}",
                "entry beta",
                r"\bibitem{alpha}",
                "entry alpha",
            ]
        ),
        encoding="utf-8",
    )

    structure = parse_aux_file(aux, bbl_path=bbl)
    assert structure is not None
    assert structure.resolve_ref("beta") == "1"
    assert structure.resolve_ref("alpha") == "2"


def test_convert_latex_to_docx_injects_real_footnotes_part(tmp_path: Path):
    output = tmp_path / "footnote.docx"
    latex = r"""
\documentclass{article}
\begin{document}
Hello\footnote{This is a real footnote.} world.
\end{document}
"""
    convert_latex_to_docx(latex_content=latex, output_path=output)
    assert output.exists()

    with zipfile.ZipFile(output, "r") as zf:
        names = set(zf.namelist())
        assert "word/footnotes.xml" in names
        footnotes_xml = zf.read("word/footnotes.xml")
        assert b"This is a real footnote." in footnotes_xml

        rels_xml = zf.read("word/_rels/document.xml.rels")
        assert b"relationships/footnotes" in rels_xml

        doc_xml = zf.read("word/document.xml")
        assert b"footnoteReference" in doc_xml


def test_convert_latex_to_docx_cite_allows_whitespace_before_optional_arg(tmp_path: Path):
    output = tmp_path / "cite_space.docx"
    latex = r"""
\documentclass{article}
\begin{document}
As shown in \cite [p.~12] {foo}, this works.
\end{document}
"""
    convert_latex_to_docx(latex_content=latex, output_path=output)
    assert output.exists()

    with zipfile.ZipFile(output, "r") as zf:
        doc_xml = zf.read("word/document.xml")
        assert b"[foo]" in doc_xml
        assert b"[]foo" not in doc_xml


def test_convert_latex_to_docx_textcite_defaults_to_bracketed_key(tmp_path: Path):
    output = tmp_path / "textcite.docx"
    latex = r"""
\documentclass{article}
\begin{document}
\textcite{foo} argues that this works.
\end{document}
"""
    convert_latex_to_docx(latex_content=latex, output_path=output)
    assert output.exists()

    with zipfile.ZipFile(output, "r") as zf:
        doc_xml = zf.read("word/document.xml")
        assert b"[foo]" in doc_xml


def test_convert_latex_to_docx_preserves_keywords_commands(tmp_path: Path):
    output = tmp_path / "keywords.docx"
    latex = r"""
\documentclass{article}
\begin{document}
\keywords{星座运维}
\KEYWORDS{Constellation Operations}
\end{document}
"""
    convert_latex_to_docx(latex_content=latex, output_path=output)
    assert output.exists()

    with zipfile.ZipFile(output, "r") as zf:
        doc_xml = zf.read("word/document.xml")
        root = ET.fromstring(doc_xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        all_text = "".join((t.text or "") for t in root.findall(".//w:t", ns))
        assert "关键词：星座运维" in all_text
        assert "Keywords: Constellation Operations" in all_text


def test_convert_latex_to_docx_uses_builtin_heading_styles_for_numbered_sections(tmp_path: Path):
    output = tmp_path / "heading_styles.docx"
    latex = r"""
\documentclass{report}
\begin{document}
\chapter{绪论}
\section{背景}
\end{document}
"""
    convert_latex_to_docx(latex_content=latex, output_path=output)
    assert output.exists()

    doc = Document(str(output))
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    assert len(non_empty) >= 2
    assert non_empty[0].style.name == "Heading 1"
    assert non_empty[1].style.name == "Heading 2"


def test_convert_latex_to_docx_uses_configurable_keyword_prefixes(tmp_path: Path):
    output = tmp_path / "keywords_configured.docx"
    latex = r"""
\documentclass{article}
\begin{document}
\keywords{A}
\KEYWORDS{B}
\end{document}
"""
    profile = DocxProfile(
        labels=LabelsConfig(
            keywords_zh_prefix="关键词-自定义:",
            keywords_en_prefix="Keywords-Custom: ",
        )
    )
    with patch("app.core.compiler.latex2docx.profile.load_profile", return_value=profile):
        convert_latex_to_docx(latex_content=latex, output_path=output, template_id="demo")

    with zipfile.ZipFile(output, "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        all_text = "".join((t.text or "") for t in root.findall(".//w:t", ns))
        assert "关键词-自定义:A" in all_text
        assert "Keywords-Custom: B" in all_text


def test_convert_latex_to_docx_uses_configurable_list_headings(tmp_path: Path):
    output = tmp_path / "list_heading_configured.docx"
    latex = r"""
\documentclass{article}
\begin{document}
\listoffigures
\listoftables
\end{document}
"""
    profile = DocxProfile(
        labels=LabelsConfig(
            list_of_figures="插图目录",
            list_of_tables="数据表目录",
        )
    )
    with patch("app.core.compiler.latex2docx.profile.load_profile", return_value=profile):
        convert_latex_to_docx(latex_content=latex, output_path=output, template_id="demo")

    doc = Document(str(output))
    text_list = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "插图目录" in text_list
    assert "数据表目录" in text_list


def test_ucas_frontmatter_uses_configurable_advisor_prefix():
    doc = Document()
    profile = DocxProfile(labels=LabelsConfig(advisor_en_prefix="Tutor: "))
    builder = UcasThesisFrontmatter(profile=profile)
    meta = WordExportMetadata(
        title="题目",
        degree="硕士",
        advisor_en="Alice",
    )

    builder._build_frontmatter(doc, meta)  # noqa: SLF001

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Tutor: Alice" in all_text
